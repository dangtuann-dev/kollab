# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_run_font(run, name='Times New Roman', size_pt=13, bold=False, italic=False):
    """Sets the font family and size for a run, ensuring Vietnamese characters render properly."""
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    
    # Force Word to apply Times New Roman to complex/non-ASCII characters
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:cs'), name)
    rPr.append(rFonts)

def remove_table_borders(table):
    """Removes all borders from a table."""
    tblPr = table._tbl.tblPr
    # Remove existing tblBorders if any
    for child in list(tblPr):
        if child.tag.endswith('tblBorders'):
            tblPr.remove(child)
    # Add empty borders
    borders = parse_xml(
        '<w:tblBorders %s>'
        '  <w:top w:val="none"/>'
        '  <w:left w:val="none"/>'
        '  <w:bottom w:val="none"/>'
        '  <w:right w:val="none"/>'
        '  <w:insideH w:val="none"/>'
        '  <w:insideV w:val="none"/>'
        '</w:tblBorders>' % nsdecls('w')
    )
    tblPr.append(borders)

def setup_section_margins(section):
    """Sets standard A4 margins (Left: 3.5cm, Right: 2.0cm, Top: 2.5cm, Bottom: 2.5cm)."""
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(2.0)

def add_body_paragraph(doc, text, bold=False, italic=False, space_after=6, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Adds a standard paragraph with 1.5 line spacing, size 13, Times New Roman."""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, 'Times New Roman', 13, bold=bold, italic=italic)
    return p

def add_bullet_paragraph(doc, text, space_after=6):
    """Adds a justified bullet paragraph styled correctly with 1.5 line spacing."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-0.63)
    
    run_prefix = p.add_run("-\t")
    set_run_font(run_prefix, 'Times New Roman', 13, bold=False)
        
    run = p.add_run(text)
    set_run_font(run, 'Times New Roman', 13, bold=False)
    return p

def fill_cell_no_border(cell, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, size_pt=13):
    """Populates a cell in an invisible layout table."""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.2
    
    # We split lines to support multi-line text or runs easily
    run = p.add_run(text)
    set_run_font(run, 'Times New Roman', size_pt, bold=bold, italic=italic)

def create_assignment():
    doc = Document()
    
    # Setup standard section margins
    section = doc.sections[0]
    setup_section_margins(section)
    
    # ==============================================================================
    # HEADER SECTION: UNIVERSITY & MOTTO (using table to place side by side)
    # ==============================================================================
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = False
    
    # Set widths: total printable width = 15.5cm. Cell 1 = 6.5cm, Cell 2 = 9.0cm.
    header_table.columns[0].cells[0].width = Cm(6.5)
    header_table.columns[1].cells[0].width = Cm(9.0)
    
    # Left Cell: School info
    cell_left = header_table.cell(0, 0)
    p_left = cell_left.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_left.paragraph_format.line_spacing = 1.15
    p_left.paragraph_format.space_after = Pt(2)
    
    run_l1 = p_left.add_run("TRƯỜNG ĐẠI HỌC NGUYỄN TẤT THÀNH\n")
    set_run_font(run_l1, 'Times New Roman', 11, bold=True)
    run_l2 = p_left.add_run("KHOA CÔNG NGHỆ THÔNG TIN\n")
    set_run_font(run_l2, 'Times New Roman', 11, bold=True)
    run_l3 = p_left.add_run("───────")
    set_run_font(run_l3, 'Times New Roman', 10, bold=False)
    
    # Right Cell: Motto info
    cell_right = header_table.cell(0, 1)
    p_right = cell_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_right.paragraph_format.line_spacing = 1.15
    p_right.paragraph_format.space_after = Pt(2)
    
    run_r1 = p_right.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n")
    set_run_font(run_r1, 'Times New Roman', 11, bold=True)
    run_r2 = p_right.add_run("Độc lập – Tự do – Hạnh phúc\n")
    set_run_font(run_r2, 'Times New Roman', 11, bold=True)
    run_r3 = p_right.add_run("────────")
    set_run_font(run_r3, 'Times New Roman', 10, bold=False)
    
    # Remove borders
    remove_table_borders(header_table)
    
    # Spacer
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(12)
    p_space.paragraph_format.space_after = Pt(12)
    
    # ==============================================================================
    # TITLE SECTION
    # ==============================================================================
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(6)
    p_title.paragraph_format.space_after = Pt(2)
    run_title = p_title.add_run("NHIỆM VỤ ĐỒ ÁN CƠ SỞ/CHUYÊN NGÀNH")
    set_run_font(run_title, 'Times New Roman', 15, bold=True)
    
    p_subtitle = doc.add_paragraph()
    p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_subtitle.paragraph_format.space_before = Pt(0)
    p_subtitle.paragraph_format.space_after = Pt(18)
    run_subtitle = p_subtitle.add_run("(Sinh viên phải đóng tờ này vào báo cáo)")
    set_run_font(run_subtitle, 'Times New Roman', 11, italic=True)
    
    # ==============================================================================
    # STUDENT & PROJECT DETAILS SECTION (using table for clean alignment)
    # ==============================================================================
    details_table = doc.add_table(rows=6, cols=2)
    details_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    details_table.autofit = False
    
    details_table.columns[0].cells[0].width = Cm(8.5)
    details_table.columns[1].cells[0].width = Cm(7.0)
    
    # Fill details
    fill_cell_no_border(details_table.cell(0, 0), "Họ và tên: TRẦN THỊ HOA")
    fill_cell_no_border(details_table.cell(0, 1), "MSSV: 2200001234")
    
    fill_cell_no_border(details_table.cell(1, 0), "Chuyên ngành: CÔNG NGHỆ THÔNG TIN")
    fill_cell_no_border(details_table.cell(1, 1), "Lớp: 22DTH1A")
    
    fill_cell_no_border(details_table.cell(2, 0), "Email: hoa.tran@student.nttu.edu.vn")
    fill_cell_no_border(details_table.cell(2, 1), "SĐT: 0909001234")
    
    # Merged rows for theme, supervisor, timeline
    row3_cell = details_table.cell(3, 0).merge(details_table.cell(3, 1))
    fill_cell_no_border(row3_cell, "Tên đề tài: Xây dựng hệ thống quản lý dự án Scrum trực quan \"Kollab\"")
    
    row4_cell = details_table.cell(4, 0).merge(details_table.cell(4, 1))
    fill_cell_no_border(row4_cell, "Giáo viên hướng dẫn: ThS. Vương Xuân Chí")
    
    row5_cell = details_table.cell(5, 0).merge(details_table.cell(5, 1))
    fill_cell_no_border(row5_cell, "Thời gian thực hiện:  15 / 03 / 2026 đến 15 / 06 / 2026")
    
    remove_table_borders(details_table)
    
    # Spacer
    p_space2 = doc.add_paragraph()
    p_space2.paragraph_format.space_before = Pt(6)
    p_space2.paragraph_format.space_after = Pt(6)
    
    # ==============================================================================
    # MÔ TẢ ĐỀ TÀI
    # ==============================================================================
    add_body_paragraph(doc, "MÔ TẢ ĐỀ TÀI:", bold=True, space_after=4)
    add_body_paragraph(
        doc,
        "Kollab là một hệ thống quản lý dự án trực quan theo mô hình Agile/Scrum, được thiết kế nhằm hỗ trợ "
        "các đội ngũ phát triển phần mềm tối ưu hóa quy trình làm việc, nâng cao khả năng cộng tác nhóm "
        "và giám sát chặt chẽ tiến độ dự án. Hệ thống giải quyết các bài toán cốt lõi của Scrum bằng cách "
        "cung cấp giao diện làm việc Single Page Application (SPA) trực quan, cho phép quản lý danh mục "
        "sản phẩm (Product Backlog), lập kế hoạch Sprint (Sprint Planning), cập nhật trạng thái tác vụ qua "
        "bảng Kanban tương tác kéo thả linh hoạt, và tự động tạo các biểu đồ đo lường hiệu suất (Burndown Chart, "
        "Velocity Chart) theo thời gian thực dựa trên các chính sách bảo mật dữ liệu nghiêm ngặt ở tầng cơ sở dữ liệu."
    )
    
    # ==============================================================================
    # NỘI DUNG VÀ PHƯƠNG PHÁP
    # ==============================================================================
    add_body_paragraph(doc, "NỘI DUNG VÀ PHƯƠNG PHÁP:", bold=True, space_after=4)
    add_bullet_paragraph(doc, "Nghiên cứu cơ sở lý thuyết về khung làm việc Scrum và phương pháp phát triển Agile.")
    add_bullet_paragraph(doc, "Phân tích yêu cầu hệ thống (chức năng, phi chức năng) và xây dựng sơ đồ ca sử dụng (Use Case).")
    add_bullet_paragraph(doc, "Thiết kế cơ sở dữ liệu quan hệ PostgreSQL chuẩn hóa với 7 bảng dữ liệu chính (profiles, projects, project_members, sprints, stories, tasks, comments).")
    add_bullet_paragraph(doc, "Lập trình giao diện Single Page Application (SPA) bằng ReactJS, TypeScript và TailwindCSS; quản lý trạng thái tập trung ở client với Zustand.")
    add_bullet_paragraph(doc, "Tích hợp dịch vụ Backend-as-a-Service (BaaS) Supabase để xác thực người dùng (Auth) và thiết lập các chính sách bảo mật mức hàng (Row Level Security - RLS) trên CSDL.")
    add_bullet_paragraph(doc, "Xây dựng bảng Kanban tương tác kéo thả bằng thư viện @dnd-kit; hiển thị biểu đồ Burndown Chart và Velocity Chart bằng thư viện Recharts.")
    add_bullet_paragraph(doc, "Kiểm thử các chức năng, đánh giá kết quả đạt được (ưu điểm, hạn chế) và đề xuất hướng phát triển tiếp theo.")
    
    # ==============================================================================
    # YÊU CẦU
    # ==============================================================================
    add_body_paragraph(doc, "YÊU CẦU:", bold=True, space_after=4)
    add_bullet_paragraph(
        doc,
        "Có kiến thức vững vàng về lập trình web hiện đại (ReactJS, TypeScript, TailwindCSS), thiết kế cơ sở dữ liệu PostgreSQL "
        "và các dịch vụ tích hợp đám mây như Supabase. Đọc hiểu tài liệu kỹ thuật chuyên ngành tiếng Anh tốt."
    )
    add_bullet_paragraph(
        doc,
        "Có tác phong làm việc khoa học, chăm chỉ, tinh thần trách nhiệm cao, có khả năng nghiên cứu độc lập cũng như cộng tác "
        "và giao tiếp hiệu quả trong môi trường làm việc nhóm."
    )
    
    # ==============================================================================
    # APPROVAL & SIGNATURES SECTION
    # ==============================================================================
    p_info_ok = doc.add_paragraph()
    p_info_ok.paragraph_format.space_before = Pt(6)
    p_info_ok.paragraph_format.space_after = Pt(12)
    run_ok = p_info_ok.add_run("Nội dung và yêu cầu đã được thông qua Bộ môn.")
    set_run_font(run_ok, 'Times New Roman', 13, italic=True)
    
    # Date
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_date.paragraph_format.space_before = Pt(6)
    p_date.paragraph_format.space_after = Pt(12)
    run_date = p_date.add_run("TP.HCM, ngày 22 tháng 06 năm 2026")
    set_run_font(run_date, 'Times New Roman', 13, italic=True)
    
    # Signature Table
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_table.autofit = False
    sig_table.columns[0].cells[0].width = Cm(7.75)
    sig_table.columns[1].cells[0].width = Cm(7.75)
    
    # Left Cell: Q. Trưởng bộ môn
    cell_dept = sig_table.cell(0, 0)
    p_dept = cell_dept.paragraphs[0]
    p_dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_dept.paragraph_format.line_spacing = 1.15
    p_dept.paragraph_format.space_after = Pt(2)
    run_dept_lbl = p_dept.add_run("Q. TRƯỞNG BỘ MÔN\n")
    set_run_font(run_dept_lbl, 'Times New Roman', 13, bold=True)
    run_dept_lbl2 = p_dept.add_run("(Ký và ghi rõ họ tên)\n\n\n\n\n\n")
    set_run_font(run_dept_lbl2, 'Times New Roman', 11, italic=True)
    run_dept_name = p_dept.add_run("ThS. Vương Xuân Chí")
    set_run_font(run_dept_name, 'Times New Roman', 13, bold=True)
    
    # Right Cell: Giảng viên hướng dẫn
    cell_sup = sig_table.cell(0, 1)
    p_sup = cell_sup.paragraphs[0]
    p_sup.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sup.paragraph_format.line_spacing = 1.15
    p_sup.paragraph_format.space_after = Pt(2)
    run_sup_lbl = p_sup.add_run("GIÁO VIÊN HƯỚNG DẪN\n")
    set_run_font(run_sup_lbl, 'Times New Roman', 13, bold=True)
    run_sup_lbl2 = p_sup.add_run("(Ký và ghi rõ họ tên)\n\n\n\n\n\n")
    set_run_font(run_sup_lbl2, 'Times New Roman', 11, italic=True)
    run_sup_name = p_sup.add_run("ThS. Vương Xuân Chí")
    set_run_font(run_sup_name, 'Times New Roman', 13, bold=True)
    
    remove_table_borders(sig_table)
    
    # Save the document
    output_filename = "Nhiem_Vu_Do_An_Kollab.docx"
    doc.save(output_filename)
    print(f"Assignment sheet successfully created and saved as {output_filename}")

if __name__ == "__main__":
    create_assignment()
