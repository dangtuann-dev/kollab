# -*- coding: utf-8 -*-
import os
import urllib.request
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def download_logo(target_path="nttu_logo.png"):
    """Downloads the official NTTU logo if it is not already available."""
    if os.path.exists(target_path):
        print(f"Logo already exists at {target_path}")
        return True
    
    url = "https://ntt.edu.vn/wp-content/uploads/2020/05/logo_ntt.png"
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
    req = urllib.request.Request(url, headers=headers)
    try:
        print(f"Downloading official NTTU logo from {url}...")
        with urllib.request.urlopen(req) as response:
            with open(target_path, "wb") as f:
                f.write(response.read())
        print("Download successful.")
        return True
    except Exception as e:
        print(f"Failed to download official logo: {e}. Fallback to generic placeholder.")
        return False

def set_run_font(run, name='Times New Roman', size_pt=13, bold=False, italic=False):
    """Sets the font family and size for a run, ensuring Vietnamese characters render properly."""
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    
    # Force Word to apply Times New Roman to complex/non-ASCII characters
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:cs'), name)
    rPr.append(rFonts)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    """Sets standard padding (margins) inside a table cell (in dxa: 20 dxa = 1 pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_border(cell, **kwargs):
    """
    Sets specific borders on a cell.
    e.g., set_cell_border(cell, bottom={"sz": 4, "val": "single", "color": "CCCCCC"})
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f'w:{edge}'
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, val in edge_data.items():
                element.set(qn(f'w:{key}'), str(val))

def setup_section_margins(section):
    """Sets standard A4 margins (Left: 3.5cm, Right: 2.0cm, Top: 2.5cm, Bottom: 2.5cm)."""
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(2.0)

def set_section_page_num_format(section, fmt, start=None):
    """Configures the page numbering format (romanLower vs decimal) and optional starting page."""
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        sectPr.append(pgNumType)
    pgNumType.set(qn('w:fmt'), fmt)
    if start is not None:
        pgNumType.set(qn('w:start'), str(start))

def add_page_number_field(run):
    """Inserts a dynamic PAGE number field code in the run."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def add_body_paragraph(doc, text, bold=False, italic=False, space_after=6, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Adds a standard justified body paragraph with 1.5 line spacing, size 13, Times New Roman."""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, 'Times New Roman', 13, bold=bold, italic=italic)
    return p

def add_bullet_paragraph(doc, text, num_prefix=None, space_after=6):
    """Adds a justified bullet/numbered list paragraph styled correctly to avoid default styles."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-0.63)
    
    if num_prefix:
        run_prefix = p.add_run(num_prefix + "\t")
        set_run_font(run_prefix, 'Times New Roman', 13, bold=True)
    else:
        run_prefix = p.add_run("•\t")
        set_run_font(run_prefix, 'Times New Roman', 13, bold=False)
        
    run = p.add_run(text)
    set_run_font(run, 'Times New Roman', 13, bold=False)
    return p

def add_custom_heading(doc, text, level, space_before=12, space_after=6):
    """Adds standard headings according to the template rules."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    
    if level == 1:
        # Chapter Heading: Center, Bold, Size 16 (We write the Chapter and Title in separate paragraphs usually,
        # but this helper creates standard large bold paragraphs too)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, 'Times New Roman', 16, bold=True)
    elif level == 2:
        # Section Heading (e.g. 1.1): Left, Bold, Size 14
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        set_run_font(run, 'Times New Roman', 14, bold=True)
    elif level == 3:
        # Subsection Heading (e.g. 1.1.1): Left, Bold, Size 13
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        set_run_font(run, 'Times New Roman', 13, bold=True)
    return p

def add_chapter_title(doc, chapter_num_text, chapter_title_text):
    """Adds a chapter heading spanning two centered bold paragraphs, as requested by the template."""
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(24)
    p1.paragraph_format.space_after = Pt(4)
    p1.paragraph_format.keep_with_next = True
    run1 = p1.add_run(chapter_num_text)
    set_run_font(run1, 'Times New Roman', 16, bold=True)
    
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(18)
    p2.paragraph_format.keep_with_next = True
    run2 = p2.add_run(chapter_title_text)
    set_run_font(run2, 'Times New Roman', 16, bold=True)

def create_styled_table(doc, rows, cols, col_widths=None):
    """Creates a centered table with repeating headers, cantSplit rows, and custom widths."""
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Prevent row splitting across pages
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement('w:cantSplit'))
    
    # Set header row to repeat on every page
    header_tr = table.rows[0]._tr.get_or_add_trPr()
    header_tr.append(OxmlElement('w:tblHeader'))
    
    # Apply column widths if specified
    if col_widths:
        for i, width in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(width)
                
    return table

def fill_cell(cell, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, size_pt=12, bg_color=None):
    """Populates a table cell, setting padding, font, and border options."""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, 'Times New Roman', size_pt, bold=bold, italic=italic)
    
    # Cell Padding (dxa)
    set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
    
    # Background color
    if bg_color:
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        
    # Standard clean thin gray border
    border_format = {"sz": 4, "val": "single", "color": "CCCCCC"}
    set_cell_border(
        cell,
        top=border_format,
        bottom=border_format,
        left=border_format,
        right=border_format
    )

def add_table_title(doc, label, title):
    """Adds a table title ABOVE the table (centered, bold label, normal title, size 12)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    
    run_lbl = p.add_run(label + ": ")
    set_run_font(run_lbl, 'Times New Roman', 12, bold=True)
    
    run_title = p.add_run(title)
    set_run_font(run_title, 'Times New Roman', 12, italic=True)
    return p

def add_figure_caption(doc, label, title, source_info=None):
    """Adds a figure caption BELOW the figure (centered, bold label, italic title, source info)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    
    run_lbl = p.add_run(label + ": ")
    set_run_font(run_lbl, 'Times New Roman', 11, bold=True)
    
    run_title = p.add_run(title)
    set_run_font(run_title, 'Times New Roman', 11, italic=True)
    
    if source_info:
        p_src = doc.add_paragraph()
        p_src.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_src.paragraph_format.space_before = Pt(0)
        p_src.paragraph_format.space_after = Pt(12)
        run_src = p_src.add_run(f"({source_info})")
        set_run_font(run_src, 'Times New Roman', 10, italic=True)
    return p

def add_toc_line(doc, title, page_str):
    """Adds a Table of Contents entry using dot leaders and right aligned page numbers."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    
    # 15.5 cm is the printable width. Set tab stop there.
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(15.5), alignment=WD_TAB_ALIGNMENT.RIGHT, leader=WD_TAB_LEADER.DOTS)
    
    # Use bold for top-level entries (Chapter, Lời Cảm Ơn, etc.)
    is_major = "CHƯƠNG" in title or "LỜI CẢM ƠN" in title or "KẾT LUẬN" in title or "PHỤ LỤC" in title or "DANH MỤC" in title or "MỤC LỤC" in title or "PHIẾU CHẤM" in title
    
    run_title = p.add_run(title + "\t")
    set_run_font(run_title, 'Times New Roman', 12 if is_major else 11, bold=is_major)
    
    run_page = p.add_run(page_str)
    set_run_font(run_page, 'Times New Roman', 12 if is_major else 11, bold=is_major)
    return p

def create_report():
    logo_path = "nttu_logo.png"
    has_logo = download_logo(logo_path)
    
    doc = Document()
    
    # ==============================================================================
    # SECTION 1: COVER & INNER COVER (No Page Numbering)
    # ==============================================================================
    sec_cover = doc.sections[0]
    setup_section_margins(sec_cover)
    # Disable headers/footers
    sec_cover.header.is_linked_to_previous = False
    sec_cover.footer.is_linked_to_previous = False
    
    # Cover Page Elements
    p_univ = doc.add_paragraph()
    p_univ.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_univ.paragraph_format.space_before = Pt(0)
    p_univ.paragraph_format.space_after = Pt(2)
    run_univ = p_univ.add_run("TRƯỜNG ĐẠI HỌC NGUYỄN TẤT THÀNH")
    set_run_font(run_univ, 'Times New Roman', 16, bold=True)
    
    p_dept = doc.add_paragraph()
    p_dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_dept.paragraph_format.space_before = Pt(0)
    p_dept.paragraph_format.space_after = Pt(18)
    run_dept = p_dept.add_run("KHOA CÔNG NGHỆ THÔNG TIN")
    set_run_font(run_dept, 'Times New Roman', 16, bold=True)
    
    # Logo placement
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_before = Pt(12)
    p_logo.paragraph_format.space_after = Pt(24)
    if has_logo:
        try:
            p_logo.add_run().add_picture(logo_path, width=Cm(3.5))
        except Exception:
            run_logo_text = p_logo.add_run("[LOGO TRƯỜNG ĐH NGUYỄN TẤT THÀNH]")
            set_run_font(run_logo_text, 'Times New Roman', 14, bold=True, italic=True)
    else:
        run_logo_text = p_logo.add_run("[LOGO TRƯỜNG ĐH NGUYỄN TẤT THÀNH]")
        set_run_font(run_logo_text, 'Times New Roman', 14, bold=True, italic=True)
        
    p_type = doc.add_paragraph()
    p_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_type.paragraph_format.space_before = Pt(12)
    p_type.paragraph_format.space_after = Pt(12)
    run_type = p_type.add_run("ĐỒ ÁN CƠ SỞ")
    set_run_font(run_type, 'Times New Roman', 18, bold=True)
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(36)
    p_title.paragraph_format.line_spacing = 1.3
    run_title = p_title.add_run("Tên đề tài:\nXÂY DỰNG HỆ THỐNG QUẢN LÝ DỰ ÁN SCRUM\nTRỰC QUAN \"KOLLAB\"")
    set_run_font(run_title, 'Times New Roman', 22, bold=True)
    
    # Spacer
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(48)
    
    # Student Details
    p_details = doc.add_paragraph()
    p_details.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_details.paragraph_format.left_indent = Cm(3.0) # Align to the center-left
    p_details.paragraph_format.space_after = Pt(4)
    p_details.paragraph_format.line_spacing = 1.3
    
    def add_detail_line(p, label, value):
        r_lbl = p.add_run(label + ": ")
        set_run_font(r_lbl, 'Times New Roman', 14, bold=True)
        r_val = p.add_run(value + "\n")
        set_run_font(r_val, 'Times New Roman', 14, bold=True)
        
    add_detail_line(p_details, "Giảng viên hướng dẫn", "ThS. NGUYỄN VĂN A")
    add_detail_line(p_details, "Sinh viên thực hiện", "TRẦN THỊ HOA")
    add_detail_line(p_details, "MSSV", "2200001234")
    add_detail_line(p_details, "Khóa", "2022")
    add_detail_line(p_details, "Ngành/Chuyên ngành", "CÔNG NGHỆ THÔNG TIN")
    
    p_loc = doc.add_paragraph()
    p_loc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_loc.paragraph_format.space_before = Pt(72)
    p_loc.paragraph_format.space_after = Pt(0)
    run_loc = p_loc.add_run("TP. HỒ CHÍ MINH, THÁNG 06 NĂM 2026")
    set_run_font(run_loc, 'Times New Roman', 13, bold=True)
    
    # Page Break for Inner Cover
    doc.add_page_break()
    
    # Inner Cover (Identical text layout as Cover, but as the "inner cover" to follow template instructions)
    p_univ2 = doc.add_paragraph()
    p_univ2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_univ2.paragraph_format.space_before = Pt(0)
    p_univ2.paragraph_format.space_after = Pt(2)
    run_univ2 = p_univ2.add_run("TRƯỜNG ĐẠI HỌC NGUYỄN TẤT THÀNH")
    set_run_font(run_univ2, 'Times New Roman', 16, bold=True)
    
    p_dept2 = doc.add_paragraph()
    p_dept2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_dept2.paragraph_format.space_before = Pt(0)
    p_dept2.paragraph_format.space_after = Pt(18)
    run_dept2 = p_dept2.add_run("KHOA CÔNG NGHỆ THÔNG TIN")
    set_run_font(run_dept2, 'Times New Roman', 16, bold=True)
    
    p_logo2 = doc.add_paragraph()
    p_logo2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo2.paragraph_format.space_before = Pt(12)
    p_logo2.paragraph_format.space_after = Pt(24)
    if has_logo:
        try:
            p_logo2.add_run().add_picture(logo_path, width=Cm(3.5))
        except Exception:
            run_logo_text2 = p_logo2.add_run("[LOGO TRƯỜNG ĐH NGUYỄN TẤT THÀNH]")
            set_run_font(run_logo_text2, 'Times New Roman', 14, bold=True, italic=True)
    else:
        run_logo_text2 = p_logo2.add_run("[LOGO TRƯỜNG ĐH NGUYỄN TẤT THÀNH]")
        set_run_font(run_logo_text2, 'Times New Roman', 14, bold=True, italic=True)
        
    p_type2 = doc.add_paragraph()
    p_type2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_type2.paragraph_format.space_before = Pt(12)
    p_type2.paragraph_format.space_after = Pt(12)
    run_type2 = p_type2.add_run("ĐỒ ÁN CƠ SỞ")
    set_run_font(run_type2, 'Times New Roman', 18, bold=True)
    
    p_title2 = doc.add_paragraph()
    p_title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title2.paragraph_format.space_before = Pt(12)
    p_title2.paragraph_format.space_after = Pt(36)
    p_title2.paragraph_format.line_spacing = 1.3
    run_title2 = p_title2.add_run("Tên đề tài:\nXÂY DỰNG HỆ THỐNG QUẢN LÝ DỰ ÁN SCRUM\nTRỰC QUAN \"KOLLAB\"")
    set_run_font(run_title2, 'Times New Roman', 22, bold=True)
    
    # Spacer
    p_space2 = doc.add_paragraph()
    p_space2.paragraph_format.space_before = Pt(48)
    
    # Student Details
    p_details2 = doc.add_paragraph()
    p_details2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_details2.paragraph_format.left_indent = Cm(3.0)
    p_details2.paragraph_format.space_after = Pt(4)
    p_details2.paragraph_format.line_spacing = 1.3
    add_detail_line(p_details2, "Giảng viên hướng dẫn", "ThS. NGUYỄN VĂN A")
    add_detail_line(p_details2, "Sinh viên thực hiện", "TRẦN THỊ HOA")
    add_detail_line(p_details2, "MSSV", "2200001234")
    add_detail_line(p_details2, "Khóa", "2022")
    add_detail_line(p_details2, "Ngành/Chuyên ngành", "CÔNG NGHỆ THÔNG TIN")
    
    p_loc2 = doc.add_paragraph()
    p_loc2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_loc2.paragraph_format.space_before = Pt(72)
    p_loc2.paragraph_format.space_after = Pt(0)
    run_loc2 = p_loc2.add_run("TP. HỒ CHÍ MINH, THÁNG 06 NĂM 2026")
    set_run_font(run_loc2, 'Times New Roman', 13, bold=True)
    
    # ==============================================================================
    # SECTION 2: FRONT MATTER (Roman Numeral Page Numbering)
    # ==============================================================================
    sec_front = doc.add_section(start_type=WD_SECTION_START.NEW_PAGE)
    setup_section_margins(sec_front)
    sec_front.footer.is_linked_to_previous = False
    
    # Configure Roman lower numbering (i, ii, iii)
    set_section_page_num_format(sec_front, "romanLower", start=1)
    
    # Add page number to Section 2 footer, centered
    f_para = sec_front.footer.paragraphs[0]
    f_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_run = f_para.add_run()
    add_page_number_field(f_run)
    set_run_font(f_run, 'Times New Roman', 11)
    
    # --- GRADING SHEET (PHIẾU CHẤM) ---
    p_t1 = doc.add_paragraph()
    p_t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_t1.paragraph_format.space_before = Pt(0)
    p_t1.paragraph_format.space_after = Pt(2)
    run_t1 = p_t1.add_run("TRƯỜNG ĐẠI HỌC NGUYỄN TẤT THÀNH")
    set_run_font(run_t1, 'Times New Roman', 12, bold=True)
    
    p_t2 = doc.add_paragraph()
    p_t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_t2.paragraph_format.space_before = Pt(0)
    p_t2.paragraph_format.space_after = Pt(2)
    run_t2 = p_t2.add_run("TRUNG TÂM KHẢO THÍ")
    set_run_font(run_t2, 'Times New Roman', 12, bold=True)
    
    p_t3 = doc.add_paragraph()
    p_t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_t3.paragraph_format.space_before = Pt(0)
    p_t3.paragraph_format.space_after = Pt(12)
    run_t3 = p_t3.add_run("KỲ THI KẾT THÚC HỌC PHẦN\nHỌC KỲ II - NĂM HỌC 2025 - 2026")
    set_run_font(run_t3, 'Times New Roman', 12, bold=True)
    
    p_t4 = doc.add_paragraph()
    p_t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_t4.paragraph_format.space_before = Pt(12)
    p_t4.paragraph_format.space_after = Pt(18)
    run_t4 = p_t4.add_run("PHIẾU CHẤM THI TIỂU LUẬN/ĐỒ ÁN")
    set_run_font(run_t4, 'Times New Roman', 14, bold=True)
    
    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_info.paragraph_format.space_after = Pt(6)
    p_info.paragraph_format.line_spacing = 1.3
    
    def add_info_run(p, label, val, newline=True):
        r_lbl = p.add_run(label + ": ")
        set_run_font(r_lbl, 'Times New Roman', 12, bold=True)
        r_val = p.add_run(val + ("\n" if newline else ""))
        set_run_font(r_val, 'Times New Roman', 12)
        
    add_info_run(p_info, "Môn thi", "Đồ án cơ sở", newline=False)
    p_info.add_run(" \t\t\t\t")
    add_info_run(p_info, "Lớp học phần", "22DTH1A", newline=True)
    add_info_run(p_info, "Nhóm sinh viên thực hiện", "Nhóm 05", newline=True)
    add_info_run(p_info, "1. Trần Thị Hoa (Trưởng nhóm)", "MSSV: 2200001234 - Đóng góp: 100%", newline=True)
    add_info_run(p_info, "Ngày thi", "25/06/2026", newline=False)
    p_info.add_run(" \t\t\t\t")
    add_info_run(p_info, "Phòng thi", "A.502", newline=True)
    add_info_run(p_info, "Đề tài tiểu luận/báo cáo", "Xây dựng hệ thống quản lý dự án Scrum trực quan \"Kollab\"", newline=True)
    
    add_body_paragraph(doc, "Phần đánh giá của giảng viên (căn cứ trên thang rubrics của môn học):", bold=True, space_after=12)
    
    # Table widths: Total printable width is 15.5 cm.
    # Col widths: Criteria (5.5 cm), Assessment (5.0 cm), Max Score (2.0 cm), Achieved Score (3.0 cm)
    table_grade = create_styled_table(doc, rows=5, cols=4, col_widths=[5.5, 5.0, 2.0, 3.0])
    
    headers = ["Tiêu chí (theo CĐR HP)", "Đánh giá của GV", "Điểm tối đa", "Điểm đạt được"]
    for i, h in enumerate(headers):
        fill_cell(table_grade.cell(0, i), h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=11, bg_color="F2F2F2")
        
    grade_data = [
        ["Cấu trúc của báo cáo (Đầy đủ phần mở đầu, các chương, kết luận, phụ lục, tài liệu tham khảo, đúng định dạng)", "", "2.0", ""],
        ["Nội dung (Cơ sở lý luận rõ ràng, phân tích thiết kế hệ thống chi tiết, cấu trúc CSDL chuẩn hóa, cài đặt hệ thống hoàn chỉnh)", "", "5.0", ""],
        ["Trình bày (Bố cục trực quan, lập luận logic, các hình vẽ/bảng biểu rõ ràng, trả lời câu hỏi vấn đáp tốt)", "", "3.0", ""],
        ["TỔNG ĐIỂM", "", "10.0", ""]
    ]
    
    for row_idx, r_data in enumerate(grade_data):
        is_total = (row_idx == len(grade_data) - 1)
        for col_idx, val in enumerate(r_data):
            cell = table_grade.cell(row_idx + 1, col_idx)
            # Center numbers
            align = WD_ALIGN_PARAGRAPH.CENTER if col_idx >= 2 else WD_ALIGN_PARAGRAPH.LEFT
            fill_cell(cell, val, bold=is_total, align=align, size_pt=11, bg_color="EAEAEA" if is_total else None)
            
    p_sign = doc.add_paragraph()
    p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_sign.paragraph_format.space_before = Pt(24)
    run_sign = p_sign.add_run("Giảng viên chấm thi\n")
    set_run_font(run_sign, 'Times New Roman', 12, bold=True, italic=False)
    run_sign_space = p_sign.add_run("(ký, ghi rõ họ tên)\n\n\n\n\n_______________________")
    set_run_font(run_sign_space, 'Times New Roman', 11, italic=True)
    
    doc.add_page_break()
    
    # --- ACKNOWLEDGMENTS (LỜI CẢM ƠN) ---
    p_ack_head = doc.add_paragraph()
    p_ack_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ack_head.paragraph_format.space_before = Pt(12)
    p_ack_head.paragraph_format.space_after = Pt(18)
    run_ack_head = p_ack_head.add_run("LỜI CẢM ƠN")
    set_run_font(run_ack_head, 'Times New Roman', 14, bold=True)
    
    add_body_paragraph(doc, "Lời đầu tiên, em xin bày tỏ lòng biết ơn sâu sắc tới các thầy, cô giáo khoa Công nghệ thông tin Trường Đại học Nguyễn Tất Thành, những người đã truyền dạy cho em những kiến thức quý báu và định hướng tư duy lập trình khoa học trong suốt thời gian học tập vừa qua.")
    add_body_paragraph(doc, "Đặc biệt, em xin gửi lời cảm ơn trân trọng nhất tới ThS. Nguyễn Văn A, người thầy đã dành nhiều thời gian, công sức để tận tình hướng dẫn, chỉ bảo và đưa ra những lời khuyên chuyên môn vô cùng giá trị, giúp em định hình giải pháp và hoàn thiện đồ án cơ sở này một cách đúng hướng.")
    add_body_paragraph(doc, "Dù đã có nhiều cố gắng trong việc nghiên cứu và thực hiện, song do giới hạn về mặt thời gian cũng như kinh nghiệm thực tiễn, đồ án chắc chắn không tránh khỏi những thiếu sót nhất định. Em rất mong nhận được những ý kiến đóng góp, nhận xét và phê bình quý báu từ quý thầy cô hội đồng để sản phẩm này ngày càng hoàn thiện hơn, mở ra những cơ hội phát triển ứng dụng trong thực tiễn.")
    add_body_paragraph(doc, "Kính chúc quý thầy cô luôn dồi dào sức khỏe, hạnh phúc và gặt hái được nhiều thành công trong sự nghiệp giáo dục cao quý!")
    
    p_ack_sign = doc.add_paragraph()
    p_ack_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_ack_sign.paragraph_format.space_before = Pt(24)
    run_ack_sign = p_ack_sign.add_run("TP. Hồ Chí Minh, ngày 22 tháng 06 năm 2026\nSinh viên thực hiện\n\n\n\nTRẦN THỊ HOA")
    set_run_font(run_ack_sign, 'Times New Roman', 12, bold=True)
    
    doc.add_page_break()
    
    # --- TABLE OF CONTENTS (MỤC LỤC) ---
    p_toc_head = doc.add_paragraph()
    p_toc_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_toc_head.paragraph_format.space_before = Pt(12)
    p_toc_head.paragraph_format.space_after = Pt(18)
    run_toc_head = p_toc_head.add_run("MỤC LỤC")
    set_run_font(run_toc_head, 'Times New Roman', 14, bold=True)
    
    add_toc_line(doc, "PHIẾU CHẤM THI TIỂU LUẬN/ĐỒ ÁN", "iii")
    add_toc_line(doc, "LỜI CẢM ƠN", "iv")
    add_toc_line(doc, "DANH MỤC CÁC BẢNG BIỂU", "vi")
    add_toc_line(doc, "DANH MỤC CÁC HÌNH VẼ", "vii")
    add_toc_line(doc, "DANH MỤC CÁC TỪ VIẾT TẮT", "viii")
    
    # We estimate the page numbers based on a standard rendering.
    add_toc_line(doc, "CHƯƠNG 1: GIỚI THIỆU ĐỀ TÀI VÀ CÔNG NGHỆ SỬ DỤNG", "1")
    add_toc_line(doc, "  1.1 Lý do chọn đề tài", "1")
    add_toc_line(doc, "  1.2 Mục tiêu đề tài", "1")
    add_toc_line(doc, "  1.3 Đối tượng và phạm vi nghiên cứu", "2")
    add_toc_line(doc, "  1.4 Công nghệ sử dụng", "2")
    add_toc_line(doc, "    1.4.1 ReactJS & TypeScript", "2")
    add_toc_line(doc, "    1.4.2 Vite Build Tool", "3")
    add_toc_line(doc, "    1.4.3 Supabase (Backend-as-a-Service)", "3")
    add_toc_line(doc, "    1.4.4 TailwindCSS & Zustand", "4")
    
    add_toc_line(doc, "CHƯƠNG 2: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", "5")
    add_toc_line(doc, "  2.1 Yêu cầu hệ thống", "5")
    add_toc_line(doc, "    2.1.1 Yêu cầu chức năng", "5")
    add_toc_line(doc, "    2.1.2 Yêu cầu phi chức năng", "6")
    add_toc_line(doc, "  2.2 Sơ đồ ca sử dụng (Use Case Diagram)", "6")
    add_toc_line(doc, "  2.3 Thiết kế cơ sở dữ liệu", "8")
    add_toc_line(doc, "    2.3.1 Danh mục các bảng", "8")
    add_toc_line(doc, "    2.3.2 Cấu trúc chi tiết các bảng dữ liệu", "8")
    
    add_toc_line(doc, "CHƯƠNG 3: KẾT QUẢ ĐẠT ĐƯỢC VÀ PHÁT TRIỂN CHI TIẾT", "12")
    add_toc_line(doc, "  3.1 Giao diện và các chức năng đã xây dựng", "12")
    add_toc_line(doc, "    3.1.1 Phân hệ Xác thực (Authentication)", "12")
    add_toc_line(doc, "    3.1.2 Quản lý dự án (Projects Module)", "13")
    add_toc_line(doc, "    3.1.3 Quản lý Backlog và User Story", "14")
    add_toc_line(doc, "    3.1.4 Bảng Kanban và Quản lý Sprint", "15")
    add_toc_line(doc, "    3.1.5 Quản lý thành viên (Members Module)", "16")
    add_toc_line(doc, "    3.1.6 Phân tích báo cáo (Reports Module)", "17")
    add_toc_line(doc, "  3.2 Đánh giá kết quả đạt được", "18")
    add_toc_line(doc, "    3.2.1 Ưu điểm", "18")
    add_toc_line(doc, "    3.2.2 Hạn chế", "18")
    
    add_toc_line(doc, "KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", "19")
    add_toc_line(doc, "  Kết luận chung", "19")
    add_toc_line(doc, "  Hướng phát triển tiếp theo", "19")
    add_toc_line(doc, "PHỤ LỤC", "20")
    add_toc_line(doc, "  Hướng dẫn cài đặt và cấu hình dự án", "20")
    add_toc_line(doc, "DANH MỤC TÀI LIỆU THAM KHẢO", "22")
    
    doc.add_page_break()
    
    # --- LIST OF TABLES (DANH MỤC BẢNG) ---
    p_lot_head = doc.add_paragraph()
    p_lot_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_lot_head.paragraph_format.space_before = Pt(12)
    p_lot_head.paragraph_format.space_after = Pt(18)
    run_lot_head = p_lot_head.add_run("DANH MỤC BẢNG")
    set_run_font(run_lot_head, 'Times New Roman', 14, bold=True)
    
    add_toc_line(doc, "Bảng 1.1: So sánh các công cụ quản lý dự án Scrum hiện nay", "2")
    add_toc_line(doc, "Bảng 2.1: Danh sách tổng hợp các bảng trong cơ sở dữ liệu", "8")
    add_toc_line(doc, "Bảng 2.2: Cấu trúc chi tiết bảng profiles", "9")
    add_toc_line(doc, "Bảng 2.3: Cấu trúc chi tiết bảng projects", "9")
    add_toc_line(doc, "Bảng 2.4: Cấu trúc chi tiết bảng project_members", "10")
    add_toc_line(doc, "Bảng 2.5: Cấu trúc chi tiết bảng sprints", "10")
    add_toc_line(doc, "Bảng 2.6: Cấu trúc chi tiết bảng user_stories", "11")
    add_toc_line(doc, "Bảng 2.7: Cấu trúc chi tiết bảng tasks", "11")
    add_toc_line(doc, "Bảng 2.8: Cấu trúc chi tiết bảng comments", "11")
    add_toc_line(doc, "Bảng 3.1: Các API Endpoint và RLS Policies của Supabase", "18")
    add_toc_line(doc, "Bảng 3.2: Thống kê số lượng dòng mã (LOC) theo các file chính", "18")
    
    doc.add_page_break()
    
    # --- LIST OF FIGURES (DANH MỤC HÌNH VẼ) ---
    p_lof_head = doc.add_paragraph()
    p_lof_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_lof_head.paragraph_format.space_before = Pt(12)
    p_lof_head.paragraph_format.space_after = Pt(18)
    run_lof_head = p_lof_head.add_run("DANH MỤC HÌNH VẼ")
    set_run_font(run_lof_head, 'Times New Roman', 14, bold=True)
    
    add_toc_line(doc, "Hình 1.1: Mô hình hoạt động cốt lõi của Scrum Framework", "1")
    add_toc_line(doc, "Hình 2.1: Sơ đồ ca sử dụng (Use Case Diagram) tổng quát", "7")
    add_toc_line(doc, "Hình 2.2: Sơ đồ thực thể quan hệ (ERD) chi tiết", "8")
    add_toc_line(doc, "Hình 3.1: Giao diện Đăng nhập và Đăng ký", "13")
    add_toc_line(doc, "Hình 3.2: Giao diện danh sách Dự án cá nhân", "14")
    add_toc_line(doc, "Hình 3.3: Giao diện Backlog quản lý User Story", "15")
    add_toc_line(doc, "Hình 3.4: Bảng Kanban kéo thả cho các thẻ nhiệm vụ trong Sprint", "16")
    add_toc_line(doc, "Hình 3.5: Giao diện quản lý thành viên dự án và vai trò", "17")
    add_toc_line(doc, "Hình 3.6: Biểu đồ báo cáo Burndown Chart và Velocity Chart", "17")
    add_toc_line(doc, "Hình 3.7: Trang cấu hình dự án", "18")
    
    doc.add_page_break()
    
    # --- LIST OF ABBREVIATIONS (DANH MỤC TỪ VIẾT TẮT) ---
    p_ab_head = doc.add_paragraph()
    p_ab_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ab_head.paragraph_format.space_before = Pt(12)
    p_ab_head.paragraph_format.space_after = Pt(18)
    run_ab_head = p_ab_head.add_run("KÍ HIỆU CÁC CỤM TỪ VIẾT TẮT")
    set_run_font(run_ab_head, 'Times New Roman', 14, bold=True)
    
    # Widths: Abbreviation (4.0 cm), Meaning (11.5 cm)
    table_ab = create_styled_table(doc, rows=12, cols=2, col_widths=[4.0, 11.5])
    
    fill_cell(table_ab.cell(0, 0), "Chữ viết tắt", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=11, bg_color="F2F2F2")
    fill_cell(table_ab.cell(0, 1), "Ý nghĩa cụ thể", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=11, bg_color="F2F2F2")
    
    ab_data = [
        ["API", "Application Programming Interface (Giao diện lập trình ứng dụng)"],
        ["CSDL", "Cơ sở dữ liệu"],
        ["CSS", "Cascading Style Sheets (Ngôn ngữ định dạng giao diện)"],
        ["DBMS", "Database Management System (Hệ quản trị cơ sở dữ liệu)"],
        ["DOM", "Document Object Model (Mô hình đối tượng tài liệu)"],
        ["HTML", "HyperText Markup Language (Ngôn ngữ đánh dấu siêu văn bản)"],
        ["PO", "Product Owner (Chủ sở hữu sản phẩm - vai trò trong Scrum)"],
        ["RLS", "Row Level Security (Bảo mật mức hàng trong cơ sở dữ liệu)"],
        ["SM", "Scrum Master (Người điều phối Scrum)"],
        ["SPA", "Single Page Application (Ứng dụng web đơn trang)"],
        ["UI/UX", "User Interface / User Experience (Giao diện / Trải nghiệm người dùng)"]
    ]
    
    for row_idx, r_data in enumerate(ab_data):
        for col_idx, val in enumerate(r_data):
            cell = table_ab.cell(row_idx + 1, col_idx)
            align = WD_ALIGN_PARAGRAPH.CENTER if col_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            fill_cell(cell, val, bold=(col_idx == 0), align=align, size_pt=11)
            
    # ==============================================================================
    # SECTION 3: BODY CHAPTERS (Arabic Page Numbering starting from 1)
    # ==============================================================================
    sec_body = doc.add_section(start_type=WD_SECTION_START.NEW_PAGE)
    setup_section_margins(sec_body)
    sec_body.footer.is_linked_to_previous = False
    
    # Configure Arabic numbering starting from 1
    set_section_page_num_format(sec_body, "decimal", start=1)
    
    # Add page number to Section 3 footer, centered
    f_para3 = sec_body.footer.paragraphs[0]
    f_para3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_run3 = f_para3.add_run()
    add_page_number_field(f_run3)
    set_run_font(f_run3, 'Times New Roman', 11)
    
    # ------------------------------------------------------------------------------
    # CHƯƠNG 1: GIỚI THIỆU ĐỀ TÀI VÀ CÔNG NGHỆ SỬ DỤNG
    # ------------------------------------------------------------------------------
    add_chapter_title(doc, "CHƯƠNG 1", "GIỚI THIỆU ĐỀ TÀI VÀ CÔNG NGHỆ SỬ DỤNG")
    
    add_custom_heading(doc, "1.1 Lý do chọn đề tài", level=2)
    add_body_paragraph(doc, "Trong kỷ nguyên chuyển đổi số hiện nay, các quy trình phát triển phần mềm đòi hỏi sự linh hoạt, tối ưu hóa thời gian và năng lực cộng tác liên tục giữa các thành viên. Phương pháp Agile nói chung và khung làm việc Scrum nói riêng đã chứng minh tính ưu việt vượt trội và trở thành tiêu chuẩn vàng tại hầu hết các tổ chức công nghệ trên thế giới. Scrum giúp các nhóm chia nhỏ dự án thành các giai đoạn ngắn (Sprint), liên tục kiểm thử, đánh giá sản phẩm và phản hồi nhanh chóng trước những thay đổi từ thị trường.")
    add_body_paragraph(doc, "Tuy nhiên, để áp dụng hiệu quả Scrum, các đội ngũ phát triển cần các công cụ hỗ trợ theo dõi tiến độ một cách trực quan. Hiện nay, các giải pháp lớn như Jira Software, Asana hay ClickUp tuy cung cấp đầy đủ chức năng nhưng lại gặp phải các hạn chế như chi phí bản quyền quá cao đối với các startup, nhóm phát triển nhỏ hoặc giao diện cấu hình quá phức tạp. Ngược lại, các công cụ đơn giản như Trello lại thiếu các chức năng đặc thù của Scrum như quản lý Sprint, điểm câu chuyện (Story Points), biểu đồ Burndown Chart hay báo cáo Velocity.")
    add_body_paragraph(doc, "Xuất phát từ nhu cầu thực tiễn đó, em quyết định thực hiện đề tài \"Xây dựng hệ thống quản lý dự án Scrum trực quan 'Kollab'\". Ứng dụng tập trung vào việc đơn giản hóa trải nghiệm người dùng nhưng vẫn đảm bảo tính đầy đủ của một công cụ Scrum chuyên nghiệp, hỗ trợ tương tác cộng tác thời gian thực với chi phí tối ưu nhất.")
    
    add_custom_heading(doc, "1.2 Mục tiêu đề tài", level=2)
    add_body_paragraph(doc, "Đồ án cơ sở này hướng tới các mục tiêu cụ thể sau:")
    add_bullet_paragraph(doc, "Nghiên cứu sâu sắc lý thuyết về phương pháp Agile và quy trình vận hành của khung làm việc Scrum.")
    add_bullet_paragraph(doc, "Thiết kế và xây dựng thành công một ứng dụng web Single Page Application (SPA) hỗ trợ đầy đủ các vai trò chính trong Scrum: Product Owner (PO), Scrum Master (SM) và Developer (Thành viên nhóm phát triển).")
    add_bullet_paragraph(doc, "Cung cấp giao diện trực quan hỗ trợ quản lý danh mục sản phẩm (Product Backlog), lập kế hoạch Sprint (Sprint Planning), bảng Kanban cập nhật trạng thái tác vụ bằng thao tác kéo thả sinh động.")
    add_bullet_paragraph(doc, "Tích hợp hệ thống biểu đồ báo cáo tự động (Burndown Chart, Velocity Chart) dựa trên dữ liệu cập nhật theo thời gian thực để đo lường năng suất làm việc của nhóm.")
    add_bullet_paragraph(doc, "Đảm bảo tính bảo mật dữ liệu dự án tối đa thông qua cấu hình phân quyền bảo mật mức hàng (Row Level Security) trực tiếp trên cơ sở dữ liệu.")
    
    add_custom_heading(doc, "1.3 Đối tượng và phạm vi nghiên cứu", level=2)
    add_body_paragraph(doc, "Đối tượng nghiên cứu: Các quy trình, tài liệu hướng dẫn về Scrum Framework của Agile Alliance; các công nghệ lập trình web hiện đại bao gồm thư viện ReactJS, kiến trúc cơ sở dữ liệu quan hệ PostgreSQL và các mô hình Backend-as-a-Service (BaaS).")
    add_body_paragraph(doc, "Phạm vi nghiên cứu: Đồ án tập trung nghiên cứu giải pháp xây dựng ứng dụng web dành cho các nhóm phát triển phần mềm nội bộ vừa và nhỏ (quy mô từ 3 đến 15 thành viên). Hệ thống tập trung tối ưu hóa các quy trình cốt lõi của Scrum mà không tích hợp các tính năng mở rộng nằm ngoài phạm vi quản lý tiến độ.")
    
    add_custom_heading(doc, "1.4 Công nghệ sử dụng", level=2)
    add_body_paragraph(doc, "Hệ thống Kollab được phát triển dựa trên sự kết hợp của các công nghệ hiện đại nhằm đảm bảo tốc độ tải trang nhanh, trải nghiệm người dùng mượt mà và khả năng mở rộng dễ dàng.")
    
    add_custom_heading(doc, "1.4.1 ReactJS & TypeScript", level=3)
    add_body_paragraph(doc, "ReactJS là thư viện JavaScript phổ biến hàng đầu thế giới để xây dựng giao diện người dùng dựa trên các thành phần (components) tái sử dụng được. Việc kết hợp TypeScript mang lại cơ chế kiểm tra kiểu dữ liệu tĩnh (static type-checking), giúp phát hiện lỗi lập trình ngay trong quá trình viết mã, tăng cường tính rõ ràng của cấu trúc mã nguồn và nâng cao hiệu suất làm việc nhóm.")
    
    add_custom_heading(doc, "1.4.2 Vite Build Tool", level=3)
    add_body_paragraph(doc, "Vite được sử dụng thay thế cho công cụ Build truyền thống như Webpack nhờ tốc độ khởi động máy chủ phát triển (dev server) và cơ chế Hot Module Replacement (HMR) cực nhanh. Vite tận dụng tối đa tính năng ES Modules nguyên bản của trình duyệt để tối ưu hóa quy trình đóng gói ứng dụng khi deploy lên môi trường sản xuất.")
    
    add_custom_heading(doc, "1.4.3 Supabase (Backend-as-a-Service)", level=3)
    add_body_paragraph(doc, "Supabase cung cấp một giải pháp Backend hoàn chỉnh dựa trên nền tảng cơ sở dữ liệu quan hệ PostgreSQL mạnh mẽ. Thay vì xây dựng API server từ đầu, hệ thống sử dụng các dịch vụ tích hợp sẵn của Supabase:")
    add_bullet_paragraph(doc, "Supabase Auth: Quản lý đăng ký, đăng nhập và cấp mã thông báo JWT an toàn.")
    add_bullet_paragraph(doc, "PostgreSQL Database: Lưu trữ dữ liệu quan hệ chặt chẽ, hỗ trợ các truy vấn phức tạp.")
    add_bullet_paragraph(doc, "Row Level Security (RLS): Cho phép cấu hình các chính sách bảo mật trực tiếp trên từng bảng, đảm bảo thành viên dự án này không thể truy cập trái phép dữ liệu của dự án khác.")
    
    add_custom_heading(doc, "1.4.4 TailwindCSS & Zustand", level=3)
    add_body_paragraph(doc, "TailwindCSS là framework CSS hướng tiện ích (utility-first) hỗ trợ thiết kế giao diện responsive linh hoạt và nhanh chóng trực tiếp trong các tệp component. Zustand được lựa chọn để quản lý trạng thái ứng dụng (state management) phía client. Đây là một thư viện nhỏ gọn, dễ sử dụng hơn Redux, giúp lưu trữ thông tin đăng nhập, danh sách dự án hiện hành và trạng thái giao diện một cách đồng bộ.")
    
    add_table_title(doc, "Bảng 1.1", "So sánh các công cụ quản lý dự án Scrum hiện nay")
    # Col widths: Tiêu chí (3.0 cm), Jira (4.0 cm), Trello (4.0 cm), Kollab (4.5 cm)
    table_comp = create_styled_table(doc, rows=5, cols=4, col_widths=[3.0, 4.0, 4.0, 4.5])
    
    comp_headers = ["Tiêu chí so sánh", "Jira Software", "Trello", "Kollab (Hệ thống đề xuất)"]
    for i, h in enumerate(comp_headers):
        fill_cell(table_comp.cell(0, i), h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=11, bg_color="F2F2F2")
        
    comp_data = [
        ["Chi phí sử dụng", "Cao (tính phí theo mỗi người dùng hàng tháng)", "Miễn phí cơ bản, giới hạn tính năng nâng cao", "Miễn phí hoàn toàn / Tự host mã nguồn"],
        ["Tính năng Scrum", "Rất đầy đủ và phức tạp, khó cấu hình", "Thiếu (phải cài thêm Power-Ups bên thứ ba)", "Đầy đủ các tính năng cốt lõi (Sprint, Backlog, Kanban)"],
        ["Trải nghiệm sử dụng", "Nặng nề, độ trễ tải trang lớn", "Mượt mà, trực quan cao", "Mượt mà, tải trang tức thì (Vite SPA)"],
        ["Độ phức tạp cài đặt", "Yêu cầu quản trị viên có chuyên môn cao", "Rất dễ sử dụng, cấu hình nhanh", "Dễ sử dụng, thiết lập dự án chỉ trong vài giây"]
    ]
    
    for row_idx, r_data in enumerate(comp_data):
        for col_idx, val in enumerate(r_data):
            cell = table_comp.cell(row_idx + 1, col_idx)
            fill_cell(cell, val, bold=(col_idx == 0), size_pt=10)
            
    doc.add_page_break()
    
    # ------------------------------------------------------------------------------
    # CHƯƠNG 2: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG
    # ------------------------------------------------------------------------------
    add_chapter_title(doc, "CHƯƠNG 2", "PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG")
    
    add_custom_heading(doc, "2.1 Yêu cầu hệ thống", level=2)
    add_body_paragraph(doc, "Việc xác định rõ ràng các yêu cầu chức năng và phi chức năng là bước tiền đề quan trọng để xây dựng cấu trúc cơ sở dữ liệu và thiết kế giao diện phù hợp với nhu cầu của các nhóm phát triển.")
    
    add_custom_heading(doc, "2.1.1 Yêu cầu chức năng", level=3)
    add_body_paragraph(doc, "Hệ thống Kollab được chia làm các phân hệ chức năng chính sau:")
    add_bullet_paragraph(doc, "Quản lý Tài khoản (Authentication): Cho phép người dùng đăng ký tài khoản mới, xác thực đăng nhập qua email/mật khẩu, khôi phục mật khẩu khi quên.")
    add_bullet_paragraph(doc, "Quản lý Dự án (Project Management): Hỗ trợ tạo mới dự án, xem danh sách dự án cá nhân sở hữu hoặc dự án được mời tham gia, lưu trữ/lưu trữ dự án.")
    add_bullet_paragraph(doc, "Quản lý Thành viên (Member Management): Người sở hữu dự án hoặc PO có thể mời thành viên tham gia qua email và phân định vai trò rõ ràng (Product Owner, Scrum Master, Developer).")
    add_bullet_paragraph(doc, "Quản lý Backlog (Backlog Management): Cho phép tạo, sửa, xóa các User Story; ước lượng độ khó bằng Story Points; phân thứ tự ưu tiên (Critical, High, Medium, Low); gán thành viên chịu trách nhiệm.")
    add_bullet_paragraph(doc, "Quản lý Sprint (Sprint Planning & Execution): Hỗ trợ tạo các chu kỳ Sprint; kéo thả các User Story từ Backlog vào Sprint hiện hành; kích hoạt kích hoạt Sprint hoạt động.")
    add_bullet_paragraph(doc, "Bảng Kanban (Sprint Board): Hiển thị trạng thái các Story theo các cột (Backlog, ToDo, In Progress, Review, Done). Người dùng có thể kéo thả Story để cập nhật trạng thái tức thời. Cho phép tạo các tác vụ con (sub-tasks) và cập nhật số giờ ước tính / thực tế.")
    
    add_custom_heading(doc, "2.1.2 Yêu cầu phi chức năng", level=3)
    add_bullet_paragraph(doc, "Tính phản hồi nhanh (Performance): Các thao tác kéo thả, chuyển trang và cập nhật dữ liệu phải diễn ra với độ trễ dưới 500ms.")
    add_bullet_paragraph(doc, "Tính tương thích (Responsive Design): Giao diện hiển thị tốt trên cả màn hình máy tính để bàn (Desktop), máy tính xách tay (Laptop) và máy tính bảng (Tablet).")
    add_bullet_paragraph(doc, "Tính bảo mật (Security): Mọi giao tiếp với cơ sở dữ liệu phải được xác thực bằng JWT, phân quyền xem và chỉnh sửa nghiêm ngặt qua Supabase RLS ở mức bảng dữ liệu để tránh rò rỉ thông tin nội bộ giữa các công ty.")
    
    add_custom_heading(doc, "2.2 Sơ đồ ca sử dụng (Use Case Diagram)", level=2)
    add_body_paragraph(doc, "Hệ thống phân định 3 vai trò chính tương tác với các ca sử dụng của hệ thống:")
    add_bullet_paragraph(doc, "Product Owner (PO): Chịu trách nhiệm tối cao về Product Backlog. PO có quyền tạo mới, chỉnh sửa độ ưu tiên của các User Story, quản lý thành viên dự án và cấu hình hệ thống.")
    add_bullet_paragraph(doc, "Scrum Master (SM): Hỗ trợ điều phối nhóm. SM có quyền lập kế hoạch Sprint, bắt đầu hoặc hoàn thành một Sprint, xem biểu đồ báo cáo hiệu năng nhóm.")
    add_bullet_paragraph(doc, "Developer (Thành viên nhóm phát triển): Có quyền xem Backlog, cập nhật trạng thái các User Story được gán trên bảng Kanban, tạo các sub-task và viết bình luận thảo luận.")
    
    # Text-based diagram placeholder or description
    p_fig_desc = doc.add_paragraph()
    p_fig_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_fig_desc.paragraph_format.space_before = Pt(6)
    p_fig_desc.paragraph_format.space_after = Pt(6)
    run_fig_desc = p_fig_desc.add_run(
        "+------------------------------------------------------------------------+\n"
        "|                             HỆ THỐNG KOLLAB                            |\n"
        "|  [Người dùng] ---> (Đăng ký / Đăng nhập)                               |\n"
        "|                                                                        |\n"
        "|  [Product Owner] ---> (Quản lý dự án) ---> (Quản lý thành viên)         |\n"
        "|                     (Quản lý Product Backlog / Thêm Story)             |\n"
        "|                                                                        |\n"
        "|  [Scrum Master] ---> (Lập kế hoạch Sprint) ---> (Quản lý Sprint)       |\n"
        "|                     (Xem biểu đồ Burndown / Velocity Chart)            |\n"
        "|                                                                        |\n"
        "|  [Developer] ---> (Nhận nhiệm vụ) ---> (Kéo thả Kanban Board)          |\n"
        "|                   (Cập nhật Sub-tasks) ---> (Bình luận thảo luận)      |\n"
        "+------------------------------------------------------------------------+"
    )
    set_run_font(run_fig_desc, 'Consolas', 10, bold=False)
    
    add_figure_caption(doc, "Hình 2.1", "Sơ đồ ca sử dụng (Use Case Diagram) tổng quát của hệ thống Kollab")
    
    add_custom_heading(doc, "2.3 Thiết kế cơ sở dữ liệu", level=2)
    add_body_paragraph(doc, "Cơ sở dữ liệu của Kollab được xây dựng trên hệ quản trị cơ sở dữ liệu PostgreSQL. Mối quan hệ giữa các bảng được thiết lập chặt chẽ nhằm phản ánh đúng nghiệp vụ Scrum.")
    
    add_custom_heading(doc, "2.3.1 Danh mục các bảng", level=3)
    add_body_paragraph(doc, "Hệ thống bao gồm 7 bảng dữ liệu quan hệ chính như sau:")
    
    add_table_title(doc, "Bảng 2.1", "Danh sách tổng hợp các bảng trong cơ sở dữ liệu")
    # Widths: 3.5, 8.5, 3.5 = 15.5 cm
    table_db = create_styled_table(doc, rows=8, cols=3, col_widths=[3.5, 8.5, 3.5])
    
    db_headers = ["Tên bảng dữ liệu", "Vai trò và chức năng chính", "Liên kết chính"]
    for i, h in enumerate(db_headers):
        fill_cell(table_db.cell(0, i), h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=11, bg_color="F2F2F2")
        
    db_data = [
        ["profiles", "Lưu trữ thông tin cá nhân của người dùng, mở rộng từ bảng auth.users của Supabase.", "auth.users (1 - 1)"],
        ["projects", "Lưu thông tin về dự án bao gồm tên, mô tả, ngày bắt đầu/kết thúc và người sở hữu.", "profiles (N - 1)"],
        ["project_members", "Bảng trung gian quản lý danh sách thành viên tham gia từng dự án và vai trò tương ứng.", "projects (N-1), profiles (N-1)"],
        ["sprints", "Quản lý các chu kỳ phát triển phần mềm (Sprint) trong một dự án.", "projects (N - 1)"],
        ["user_stories", "Lưu trữ các User Story (nhiệm vụ lớn) thuộc Backlog hoặc được gán vào Sprint.", "projects, sprints, profiles"],
        ["tasks", "Lưu trữ các tác vụ con (sub-tasks) cần thực hiện để hoàn thành một User Story.", "user_stories (N - 1)"],
        ["comments", "Lưu trữ các ý kiến đóng góp, thảo luận của các thành viên về một User Story.", "user_stories, profiles"]
    ]
    
    for row_idx, r_data in enumerate(db_data):
        for col_idx, val in enumerate(r_data):
            cell = table_db.cell(row_idx + 1, col_idx)
            fill_cell(cell, val, bold=(col_idx == 0), size_pt=10)

    # Insert ERD section and figure caption
    add_body_paragraph(doc, "Dưới đây là Sơ đồ thực thể quan hệ (ERD) chi tiết mô tả cấu trúc liên kết cơ sở dữ liệu của hệ thống Kollab:")
    p_erd_desc = doc.add_paragraph()
    p_erd_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_erd_desc.paragraph_format.space_before = Pt(6)
    p_erd_desc.paragraph_format.space_after = Pt(6)
    run_erd_desc = p_erd_desc.add_run(
        "                +-------------------+ \n"
        "                |     profiles      | \n"
        "                +-------------------+ \n"
        "                | PK | id           |<--------+ \n"
        "                |    | full_name    |         | \n"
        "                |    | email        |         | (1) \n"
        "                |    | avatar_url   |         | \n"
        "                +-------------------+         | \n"
        "                          |                   | \n"
        "                          | (1)               | \n"
        "                          |                   | \n"
        "                          v (N)               | \n"
        "                +-------------------+         | \n"
        "                |     projects      |         | \n"
        "                +-------------------+         | \n"
        "                | PK | id           |         | \n"
        "                | FK | owner_id     |         | \n"
        "                +-------------------+         | \n"
        "                  |               |           | \n"
        "                  | (1)           | (1)       | \n"
        "                  |               |           | \n"
        "                  v (N)           v (N)       | \n"
        "         +------------+      +------------+   | \n"
        "         |  sprints   |      | prj_members|   | \n"
        "         +------------+      +------------+   | \n"
        "         | PK | id    |      | PK | id    |   | \n"
        "         | FK | prj_id|      | FK | prj_id|   | \n"
        "         +------------+      | FK | user_id|---+ \n"
        "               |             +------------+ \n"
        "               | (0..1) \n"
        "               | \n"
        "               v (N) \n"
        "         +--------------------------------+ \n"
        "         |          user_stories          | \n"
        "         +--------------------------------+ \n"
        "         | PK | id                        | \n"
        "         | FK | project_id                | \n"
        "         | FK | sprint_id                 | \n"
        "         |    | story_points              | \n"
        "         |    | priority                  | \n"
        "         |    | status (backlog/sprint/dn)| \n"
        "         | FK | assignee_id               | \n"
        "         | FK | reporter_id               | \n"
        "         +--------------------------------+ \n"
        "                  |                      | \n"
        "                  | (1)                  | (1) \n"
        "                  |                      | \n"
        "                  v (N)                  v (N) \n"
        "         +------------------+  +------------------+ \n"
        "         |      tasks       |  |     comments     | \n"
        "         +------------------+  +------------------+ \n"
        "         | PK | id          |  | PK | id          | \n"
        "         | FK | user_story_i|  | FK | user_story_i| \n"
        "         |    | status      |  | FK | user_id      | \n"
        "         | FK | assignee_id |  |    | content     | \n"
        "         |    | estimate_hrs|  +------------------+ \n"
        "         +------------------+ \n"
    )
    set_run_font(run_erd_desc, 'Consolas', 9, bold=False)
    add_figure_caption(doc, "Hình 2.2", "Sơ đồ thực thể quan hệ (ERD) chi tiết của hệ thống Kollab")
            
    add_custom_heading(doc, "2.3.2 Cấu trúc chi tiết các bảng dữ liệu", level=3)
    add_body_paragraph(doc, "Dưới đây là đặc tả chi tiết của từng bảng dữ liệu được trích xuất từ cấu trúc thiết kế cơ sở dữ liệu vật lý:")
    
    # 2.2 profiles
    add_table_title(doc, "Bảng 2.2", "Cấu trúc chi tiết bảng profiles")
    table_p = create_styled_table(doc, rows=8, cols=5, col_widths=[3.0, 3.0, 3.0, 1.5, 5.0])
    struct_headers = ["Tên trường", "Kiểu dữ liệu", "Ràng buộc", "Khóa", "Mô tả"]
    for i, h in enumerate(struct_headers):
        fill_cell(table_p.cell(0, i), h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=11, bg_color="F2F2F2")
    p_rows = [
        ["id", "uuid", "PRIMARY KEY", "PK", "Liên kết trực tiếp với bảng auth.users"],
        ["full_name", "text", "NOT NULL", "", "Họ và tên đầy đủ của người dùng"],
        ["email", "text", "UNIQUE", "", "Địa chỉ hòm thư điện tử đăng ký"],
        ["avatar_url", "text", "NULL", "", "Đường dẫn ảnh đại diện người dùng"],
        ["bio", "text", "NULL", "", "Tiểu sử cá nhân ngắn gọn"],
        ["created_at", "timestamptz", "DEFAULT now()", "", "Thời điểm tạo tài khoản"],
        ["updated_at", "timestamptz", "DEFAULT now()", "", "Thời điểm cập nhật thông tin gần nhất"]
    ]
    for row_idx, r_data in enumerate(p_rows):
        for col_idx, val in enumerate(r_data):
            fill_cell(table_p.cell(row_idx + 1, col_idx), val, bold=(col_idx == 0), size_pt=9)
            
    # 2.3 projects
    add_table_title(doc, "Bảng 2.3", "Cấu trúc chi tiết bảng projects")
    table_proj = create_styled_table(doc, rows=10, cols=5, col_widths=[3.0, 3.0, 3.0, 1.5, 5.0])
    for i, h in enumerate(struct_headers):
        fill_cell(table_proj.cell(0, i), h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=11, bg_color="F2F2F2")
    proj_rows = [
        ["id", "uuid", "PRIMARY KEY", "PK", "Mã định danh duy nhất của dự án"],
        ["name", "text", "NOT NULL", "", "Tên dự án"],
        ["description", "text", "NULL", "", "Mô tả mục tiêu dự án"],
        ["owner_id", "uuid", "REFERENCES profiles(id)", "FK", "Mã người tạo dự án"],
        ["status", "text", "CHECK (active, archived)", "", "Trạng thái hoạt động của dự án"],
        ["start_date", "date", "NULL", "", "Ngày bắt đầu dự án"],
        ["end_date", "date", "NULL", "", "Ngày kết thúc dự kiến"],
        ["created_at", "timestamptz", "DEFAULT now()", "", "Ngày tạo dự án"],
        ["updated_at", "timestamptz", "DEFAULT now()", "", "Ngày cập nhật dự án"]
    ]
    for row_idx, r_data in enumerate(proj_rows):
        for col_idx, val in enumerate(r_data):
            fill_cell(table_proj.cell(row_idx + 1, col_idx), val, bold=(col_idx == 0), size_pt=9)

    # 2.4 project_members
    add_table_title(doc, "Bảng 2.4", "Cấu trúc chi tiết bảng project_members")
    table_mem = create_styled_table(doc, rows=6, cols=5, col_widths=[3.0, 3.0, 3.0, 1.5, 5.0])
    for i, h in enumerate(struct_headers):
        fill_cell(table_mem.cell(0, i), h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=11, bg_color="F2F2F2")
    mem_rows = [
        ["id", "uuid", "PRIMARY KEY", "PK", "Mã bản ghi thành viên"],
        ["project_id", "uuid", "REFERENCES projects(id)", "FK", "Dự án tham gia"],
        ["user_id", "uuid", "REFERENCES profiles(id)", "FK", "Thành viên tham gia"],
        ["role", "text", "CHECK (product_owner, scrum_master, developer)", "", "Vai trò cụ thể trong dự án"],
        ["joined_at", "timestamptz", "DEFAULT now()", "", "Ngày tham gia dự án"]
    ]
    for row_idx, r_data in enumerate(mem_rows):
        for col_idx, val in enumerate(r_data):
            fill_cell(table_mem.cell(row_idx + 1, col_idx), val, bold=(col_idx == 0), size_pt=9)

    # 2.5 sprints
    add_table_title(doc, "Bảng 2.5", "Cấu trúc chi tiết bảng sprints")
    table_spr = create_styled_table(doc, rows=10, cols=5, col_widths=[3.0, 3.0, 3.0, 1.5, 5.0])
    for i, h in enumerate(struct_headers):
        fill_cell(table_spr.cell(0, i), h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=11, bg_color="F2F2F2")
    spr_rows = [
        ["id", "uuid", "PRIMARY KEY", "PK", "Mã định danh duy nhất của Sprint"],
        ["project_id", "uuid", "REFERENCES projects(id)", "FK", "Thuộc dự án nào"],
        ["name", "text", "NOT NULL", "", "Tên Sprint (ví dụ: Sprint 1)"],
        ["goal", "text", "NULL", "", "Mục tiêu cần đạt được trong Sprint"],
        ["status", "text", "CHECK (planning, active, completed)", "", "Trạng thái chu kỳ Sprint"],
        ["start_date", "date", "NULL", "", "Ngày bắt đầu Sprint"],
        ["end_date", "date", "NULL", "", "Ngày kết thúc Sprint"],
        ["velocity", "integer", "NULL", "", "Số lượng Story Points hoàn thành thực tế"],
        ["created_at", "timestamptz", "DEFAULT now()", "", "Thời điểm khởi tạo bản ghi"]
    ]
    for row_idx, r_data in enumerate(spr_rows):
        for col_idx, val in enumerate(r_data):
            fill_cell(table_spr.cell(row_idx + 1, col_idx), val, bold=(col_idx == 0), size_pt=9)

    # 2.6 user_stories
    add_table_title(doc, "Bảng 2.6", "Cấu trúc chi tiết bảng user_stories")
    table_story = create_styled_table(doc, rows=14, cols=5, col_widths=[3.0, 3.0, 3.0, 1.5, 5.0])
    for i, h in enumerate(struct_headers):
        fill_cell(table_story.cell(0, i), h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=11, bg_color="F2F2F2")
    story_rows = [
        ["id", "uuid", "PRIMARY KEY", "PK", "Mã bản ghi User Story"],
        ["project_id", "uuid", "REFERENCES projects(id)", "FK", "Thuộc dự án nào"],
        ["sprint_id", "uuid", "REFERENCES sprints(id)", "FK", "Được phân bổ vào Sprint nào (nếu có)"],
        ["title", "text", "NOT NULL", "", "Tiêu đề của User Story"],
        ["description", "text", "NULL", "", "Nội dung đặc tả User Story"],
        ["acceptance_criteria", "text", "NULL", "", "Tiêu chí nghiệm thu tính năng"],
        ["story_points", "integer", "CHECK (>= 0)", "", "Điểm ước lượng độ khó"],
        ["priority", "text", "CHECK (critical, high, medium, low)", "", "Độ ưu tiên giải quyết nhiệm vụ"],
        ["status", "text", "CHECK (backlog, sprint, done)", "", "Trạng thái triển khai câu chuyện"],
        ["assignee_id", "uuid", "REFERENCES profiles(id)", "FK", "Thành viên chịu trách nhiệm thực hiện"],
        ["reporter_id", "uuid", "REFERENCES profiles(id)", "FK", "Người tạo ra yêu cầu nhiệm vụ này"],
        ["order_index", "integer", "DEFAULT 0", "", "Vị trí sắp xếp thứ tự trong Backlog"],
        ["created_at", "timestamptz", "DEFAULT now()", "", "Thời điểm tạo bản ghi"]
    ]
    for row_idx, r_data in enumerate(story_rows):
        for col_idx, val in enumerate(r_data):
            fill_cell(table_story.cell(row_idx + 1, col_idx), val, bold=(col_idx == 0), size_pt=9)

    # 2.7 tasks
    add_table_title(doc, "Bảng 2.7", "Cấu trúc chi tiết bảng tasks")
    table_task = create_styled_table(doc, rows=10, cols=5, col_widths=[3.0, 3.0, 3.0, 1.5, 5.0])
    for i, h in enumerate(struct_headers):
        fill_cell(table_task.cell(0, i), h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=11, bg_color="F2F2F2")
    task_rows = [
        ["id", "uuid", "PRIMARY KEY", "PK", "Mã bản ghi tác vụ con"],
        ["user_story_id", "uuid", "REFERENCES user_stories(id)", "FK", "Thuộc User Story lớn nào"],
        ["title", "text", "NOT NULL", "", "Tiêu đề tác vụ con"],
        ["description", "text", "NULL", "", "Mô tả chi tiết cách xử lý"],
        ["status", "text", "CHECK (todo, in_progress, done)", "", "Trạng thái thực hiện tác vụ"],
        ["assignee_id", "uuid", "REFERENCES profiles(id)", "FK", "Người nhận thực hiện tác vụ"],
        ["estimate_hours", "numeric(5,2)", "DEFAULT 0.00", "", "Số giờ ước tính để hoàn thành"],
        ["actual_hours", "numeric(5,2)", "DEFAULT 0.00", "", "Số giờ thực tế đã bỏ ra"],
        ["created_at", "timestamptz", "DEFAULT now()", "", "Thời điểm khởi tạo tác vụ"]
    ]
    for row_idx, r_data in enumerate(task_rows):
        for col_idx, val in enumerate(r_data):
            fill_cell(table_task.cell(row_idx + 1, col_idx), val, bold=(col_idx == 0), size_pt=9)

    # 2.8 comments
    add_table_title(doc, "Bảng 2.8", "Cấu trúc chi tiết bảng comments")
    table_cmt = create_styled_table(doc, rows=7, cols=5, col_widths=[3.0, 3.0, 3.0, 1.5, 5.0])
    for i, h in enumerate(struct_headers):
        fill_cell(table_cmt.cell(0, i), h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=11, bg_color="F2F2F2")
    cmt_rows = [
        ["id", "uuid", "PRIMARY KEY", "PK", "Mã bản ghi bình luận"],
        ["user_story_id", "uuid", "REFERENCES user_stories(id)", "FK", "Nội dung bình luận thuộc User Story nào"],
        ["user_id", "uuid", "REFERENCES profiles(id)", "FK", "Người viết bình luận"],
        ["content", "text", "NOT NULL", "", "Nội dung bình luận"],
        ["created_at", "timestamptz", "DEFAULT now()", "", "Thời điểm đăng bình luận"],
        ["updated_at", "timestamptz", "DEFAULT now()", "", "Thời điểm sửa đổi bình luận"]
    ]
    for row_idx, r_data in enumerate(cmt_rows):
        for col_idx, val in enumerate(r_data):
            fill_cell(table_cmt.cell(row_idx + 1, col_idx), val, bold=(col_idx == 0), size_pt=9)

    doc.add_page_break()
    
    # ------------------------------------------------------------------------------
    # CHƯƠNG 3: KẾT QUẢ ĐẠT ĐƯỢC VÀ PHÁT TRIỂN CHI TIẾT
    # ------------------------------------------------------------------------------
    add_chapter_title(doc, "CHƯƠNG 3", "KẾT QUẢ ĐẠT ĐƯỢC VÀ PHÁT TRIỂN CHI TIẾT")
    
    add_custom_heading(doc, "3.1 Giao diện và các chức năng đã xây dựng", level=2)
    add_body_paragraph(doc, "Dựa trên thiết kế cơ sở dữ liệu và yêu cầu hệ thống ở chương 2, toàn bộ ứng dụng Kollab đã được lập trình hoàn thiện các mô-đun quan trọng. Dưới đây là phân tích chi tiết mã nguồn và giao diện người dùng tương ứng.")
    
    add_custom_heading(doc, "3.1.1 Phân hệ Xác thực (Authentication)", level=3)
    add_body_paragraph(doc, "Phân hệ xác thực chịu trách nhiệm bảo vệ toàn bộ trang Dashboard chính của ứng dụng bằng cơ chế Private Route. Phân hệ bao gồm các trang giao diện chính:")
    add_bullet_paragraph(doc, "LoginPage.tsx: Chứa form nhập email và mật khẩu với các luật kiểm tra nhập liệu nghiêm ngặt sử dụng thư viện react-hook-form kết hợp zod resolver. Khi đăng nhập thành công, token được lưu và trạng thái authStore (Zustand) được thiết lập.")
    add_bullet_paragraph(doc, "RegisterPage.tsx: Hỗ trợ người dùng tạo tài khoản mới. Khi đăng ký, một trigger tự động trên Supabase (handle_new_user) sẽ tạo bản ghi đồng bộ tương ứng ở bảng public.profiles.")
    add_bullet_paragraph(doc, "ForgotPasswordPage.tsx: Gửi đường dẫn đặt lại mật khẩu về hòm thư điện tử cá nhân của người dùng.")
    
    add_body_paragraph(doc, "[Hình ảnh mô phỏng: Giao diện Đăng nhập với thiết kế Dark Mode cao cấp, có các hiệu ứng hover mượt mà và nút đăng nhập thiết kế Gradient].", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_figure_caption(doc, "Hình 3.1", "Giao diện Đăng nhập và Đăng ký của hệ thống Kollab")
    
    add_custom_heading(doc, "3.1.2 Quản lý dự án (Projects Module)", level=3)
    add_body_paragraph(doc, "Đây là màn hình đầu tiên sau khi người dùng đăng nhập thành công:")
    add_bullet_paragraph(doc, "ProjectsPage.tsx: Truy xuất danh sách dự án từ cơ sở dữ liệu thông qua react-query để tự động caching và tối ưu hóa tải trọng mạng. Trang hiển thị cả hai nhóm dự án: Dự án do người dùng sở hữu và Dự án người dùng tham gia với tư cách thành viên.")
    add_bullet_paragraph(doc, "ProjectCard.tsx: Component hiển thị tóm tắt thông tin của từng dự án bao gồm: Tên dự án, mô tả ngắn, vai trò của người dùng hiện hành, số lượng thành viên đang tham gia và tiến độ chung.")
    add_bullet_paragraph(doc, "CreateProjectModal.tsx: Modal pop-up giúp người dùng tạo dự án mới nhanh chóng. Người tạo dự án sẽ mặc định có vai trò là Owner kiêm Product Owner.")
    
    add_body_paragraph(doc, "[Hình ảnh mô phỏng: Trang danh sách dự án hiển thị dạng lưới (Grid Layout) responsive, các card dự án có viền bóng đổ dịu nhẹ và hiệu ứng zoom nhẹ khi di chuột qua].", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_figure_caption(doc, "Hình 3.2", "Giao diện danh sách Dự án cá nhân và Dự án tham gia")
    
    add_custom_heading(doc, "3.1.3 Quản lý Backlog và User Story", level=3)
    add_body_paragraph(doc, "Mô-đun Backlog là trung tâm quản lý yêu cầu của Product Owner:")
    add_bullet_paragraph(doc, "BacklogPage.tsx: Quản lý danh mục các User Story chưa phân bổ và các Sprint đang lập kế hoạch. Người dùng có thể kéo thả các User Story trực tiếp từ danh mục Backlog chung vào các Sprint cụ thể nhờ thư viện @dnd-kit.")
    add_bullet_paragraph(doc, "CreateStoryModal.tsx: Cung cấp form tạo mới các User Story với các trường thông tin: Tiêu đề, mô tả chi tiết, tiêu chí nghiệm thu (Acceptance Criteria), điểm Story Points và mức độ ưu tiên.")
    add_bullet_paragraph(doc, "StoryDetailPanel.tsx: Bảng điều khiển trượt từ bên phải màn hình hiển thị toàn bộ chi tiết của một Story được chọn. Tại đây thành viên có thể gán người thực hiện, cập nhật điểm số, tạo danh sách tác vụ con (sub-tasks) và gửi bình luận thảo luận thời gian thực.")
    
    add_body_paragraph(doc, "[Hình ảnh mô phỏng: Trang quản lý Backlog phân tách rõ ràng giữa khu vực Backlog chung và danh sách Sprint, hỗ trợ thao tác kéo thả mượt mà kèm theo các huy hiệu (badges) màu sắc biểu thị độ ưu tiên].", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_figure_caption(doc, "Hình 3.3", "Giao diện Backlog quản lý các User Story và điểm số")
    
    add_custom_heading(doc, "3.1.4 Bảng Kanban và Quản lý Sprint", level=3)
    add_body_paragraph(doc, "Bảng Kanban là công cụ theo dõi tiến trình làm việc hàng ngày của nhóm phát triển trong một Sprint hoạt động:")
    add_bullet_paragraph(doc, "SprintBoardPage.tsx: Tải thông tin của Sprint đang kích hoạt (Active Sprint). Nếu không có Sprint nào hoạt động, màn hình sẽ hiển thị trạng thái trống và gợi ý người dùng lập kế hoạch tại trang Backlog.")
    add_bullet_paragraph(doc, "KanbanBoard.tsx: Chia làm 5 cột trạng thái: Backlog, ToDo, In Progress, Review và Done. Các thành viên kéo thả các thẻ Story giữa các cột để cập nhật trạng thái làm việc.")
    add_bullet_paragraph(doc, "TaskCard.tsx: Thẻ hiển thị tóm tắt của từng Story trên bảng Kanban: Tiêu đề, điểm Story Points, ảnh đại diện thành viên chịu trách nhiệm, và tiến độ hoàn thành các sub-tasks tương ứng.")
    
    add_body_paragraph(doc, "[Hình ảnh mô phỏng: Bảng Kanban 5 cột tiêu chuẩn Scrum với các thẻ nhiệm vụ trực quan, hiển thị tiến độ hoàn thành sub-task dạng thanh tiến trình (progress bar) và hỗ trợ kéo thả mượt mà].", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_figure_caption(doc, "Hình 3.4", "Bảng Kanban kéo thả cho các thẻ nhiệm vụ trong Sprint")
    
    add_custom_heading(doc, "3.1.5 Quản lý thành viên (Members Module)", level=3)
    add_body_paragraph(doc, "Trang quản lý nhân sự giúp gắn kết đội ngũ phát triển dự án:")
    add_bullet_paragraph(doc, "MembersPage.tsx: Hiển thị danh sách các thành viên hiện tại của dự án cùng thông tin liên hệ và vai trò cụ thể.")
    add_bullet_paragraph(doc, "InviteMemberModal.tsx: Hỗ trợ mời thêm thành viên vào dự án thông qua email của họ đã đăng ký trên hệ thống Kollab.")
    add_bullet_paragraph(doc, "MemberCard.tsx: Card hiển thị chi tiết avatar, họ tên, email, ngày tham gia và menu thả xuống (dropdown) để thay đổi vai trò (PO, SM, Developer) trực tiếp đối với những người dùng có quyền quản trị dự án.")
    
    add_body_paragraph(doc, "[Hình ảnh mô phỏng: Danh sách thành viên kèm avatar tròn đẹp mắt, nút gán quyền cho phép chọn nhanh vai trò thông qua trình đơn thả xuống được thiết kế tinh tế].", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_figure_caption(doc, "Hình 3.5", "Giao diện quản lý thành viên dự án và vai trò")
    
    add_custom_heading(doc, "3.1.6 Phân tích báo cáo (Reports Module)", level=3)
    add_body_paragraph(doc, "Báo cáo là công cụ quan trọng để Scrum Master đánh giá năng lực của nhóm:")
    add_bullet_paragraph(doc, "ReportsPage.tsx: Tích hợp các biểu đồ phân tích trực quan sử dụng thư viện recharts.")
    add_bullet_paragraph(doc, "BurndownChart.tsx: Biểu đồ giải phóng công việc hiển thị đường cơ sở lý thuyết (đường chéo màu xám) và đường thực tế (màu đỏ) biểu thị tổng số điểm Story Points còn lại theo từng ngày của Sprint. Giúp phát hiện nhanh nhóm đang bị trễ tiến độ hay hoàn thành sớm.")
    add_bullet_paragraph(doc, "VelocityChart.tsx: Biểu đồ cột so sánh tổng số điểm Story Points cam kết khi bắt đầu Sprint và số điểm thực tế đã hoàn thành qua các chu kỳ Sprint gần nhất. Hỗ trợ nhóm ước lượng chính xác hơn khối lượng công việc cho các Sprint tiếp theo.")
    
    add_body_paragraph(doc, "[Hình ảnh mô phỏng: Đồ thị Burndown Chart sắc nét, hiển thị rõ ràng đường xu hướng lý thuyết và các điểm ghi nhận thực tế của dự án].", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_figure_caption(doc, "Hình 3.6", "Biểu đồ báo cáo Burndown Chart và Velocity Chart")
    
    # 3.1.7 Settings page
    add_custom_heading(doc, "3.1.7 Cấu hình dự án (Settings)", level=3)
    add_body_paragraph(doc, "Trang ProjectSettingsPage.tsx cho phép chủ sở hữu dự án cấu hình lại các thông tin cơ bản của dự án như tên gọi, mô tả, ngày bắt đầu/kết thúc hoặc thực hiện lưu trữ (archive) dự án khi đã hoàn thành toàn bộ sản phẩm.")
    
    add_body_paragraph(doc, "[Hình ảnh mô phỏng: Trang cấu hình dự án với các trường nhập liệu cân đối, có nút Lưu thay đổi nổi bật và nút Xóa/Lưu trữ dự án màu đỏ cảnh báo].", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_figure_caption(doc, "Hình 3.7", "Trang cấu hình dự án")
    
    add_custom_heading(doc, "3.2 Đánh giá kết quả đạt được", level=2)
    add_body_paragraph(doc, "Trải qua quá trình triển khai, ứng dụng Kollab đã cơ bản đạt được các mục tiêu thiết lập từ ban đầu. Dưới đây là bảng tổng hợp đánh giá bảo mật cơ sở dữ liệu qua các chính sách RLS thực thi trực tiếp trên hệ thống:")
    
    add_table_title(doc, "Bảng 3.1", "Các API Endpoint và RLS Policies của Supabase")
    # Widths: 3.0, 6.0, 6.5 = 15.5 cm
    table_api = create_styled_table(doc, rows=8, cols=3, col_widths=[3.0, 6.0, 6.5])
    
    api_headers = ["Bảng dữ liệu", "Loại chính sách RLS", "Mô tả cơ chế phân quyền thực tế"]
    for i, h in enumerate(api_headers):
        fill_cell(table_api.cell(0, i), h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=11, bg_color="F2F2F2")
        
    api_data = [
        ["profiles", "Đọc công khai / Sửa đổi cá nhân", "Cho phép mọi thành viên đọc hồ sơ cá nhân của nhau. Chỉ chính chủ mới có quyền UPDATE hồ sơ."],
        ["projects", "Đọc theo thành viên / Sửa bởi chủ sở hữu", "Chỉ những thành viên thuộc dự án mới có quyền SELECT dự án. Chỉ chủ sở hữu (Owner) mới có quyền UPDATE/DELETE."],
        ["project_members", "Đọc theo thành viên / Sửa bởi PO", "Thành viên dự án được quyền xem danh sách nhóm. Chỉ PO hoặc Owner mới được quyền INSERT/UPDATE/DELETE thành viên."],
        ["sprints", "Đọc theo thành viên / Sửa bởi PO và SM", "Thành viên dự án được quyền xem các Sprint. Chỉ PO và SM mới được toàn quyền quản lý Sprint (INSERT/UPDATE/DELETE)."],
        ["user_stories", "Đọc và sửa bởi thành viên / Tạo bởi PO", "Mọi thành viên được xem và sửa thông tin User Story. Chỉ PO mới được tạo (INSERT) hoặc xóa (DELETE) User Story."],
        ["tasks", "Toàn quyền quản lý bởi thành viên", "Mọi thành viên trong dự án đều có quyền tạo, sửa, xóa các tác vụ con (sub-tasks) thuộc dự án đó."],
        ["comments", "Đọc theo thành viên / Sửa bởi tác giả", "Thành viên được xem bình luận. Chỉ tác giả bình luận (user_id) mới có quyền UPDATE. Tác giả hoặc PO có quyền DELETE bình luận."]
    ]
    
    for row_idx, r_data in enumerate(api_data):
        for col_idx, val in enumerate(r_data):
            cell = table_api.cell(row_idx + 1, col_idx)
            fill_cell(cell, val, bold=(col_idx == 0), size_pt=10)
            
    add_body_paragraph(doc, "Thống kê quy mô mã nguồn triển khai thực tế của dự án:")
    add_table_title(doc, "Bảng 3.2", "Thống kê số lượng dòng mã (LOC) theo các file chính")
    # Widths: 4.0, 6.5, 5.0 = 15.5 cm
    table_loc = create_styled_table(doc, rows=7, cols=3, col_widths=[4.0, 6.5, 5.0])
    
    loc_headers = ["Phân hệ / Module", "Các tệp mã nguồn chính", "Số dòng mã ước tính (LOC)"]
    for i, h in enumerate(loc_headers):
        fill_cell(table_loc.cell(0, i), h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=11, bg_color="F2F2F2")
        
    loc_data = [
        ["Authentication (Xác thực)", "LoginPage.tsx, RegisterPage.tsx, useAuth.ts, authStore.ts", "650 dòng"],
        ["Projects (Dự án)", "ProjectsPage.tsx, ProjectCard.tsx, CreateProjectModal.tsx, projectStore.ts", "450 dòng"],
        ["Backlog (Danh mục story)", "BacklogPage.tsx, CreateStoryModal.tsx, StoryDetailPanel.tsx, StoryCard.tsx", "1100 dòng"],
        ["Sprint & Kanban (Thực thi)", "SprintBoardPage.tsx, KanbanBoard.tsx, TaskCard.tsx, SprintHeader.tsx", "850 dòng"],
        ["Members & Reports (Nhân sự & Báo cáo)", "MembersPage.tsx, ReportsPage.tsx, BurndownChart.tsx, VelocityChart.tsx", "750 dòng"],
        ["Database Schema (Cơ sở dữ liệu)", "supabase_schema.sql", "330 dòng"]
    ]
    
    for row_idx, r_data in enumerate(loc_data):
        for col_idx, val in enumerate(r_data):
            cell = table_loc.cell(row_idx + 1, col_idx)
            align = WD_ALIGN_PARAGRAPH.CENTER if col_idx == 2 else WD_ALIGN_PARAGRAPH.LEFT
            fill_cell(cell, val, bold=(col_idx == 0), align=align, size_pt=10)

    add_custom_heading(doc, "3.2.1 Ưu điểm", level=3)
    add_bullet_paragraph(doc, "Hiệu năng tải trang cực nhanh nhờ cấu trúc Single Page Application chạy trên nền Vite và ReactJS.")
    add_bullet_paragraph(doc, "Thao tác cập nhật trạng thái nhiệm vụ bằng kéo thả trực quan giúp nâng cao đáng kể trải nghiệm người dùng so với việc chọn menu thả xuống truyền thống.")
    add_bullet_paragraph(doc, "Hệ thống phân quyền RLS chặt chẽ trực tiếp dưới tầng CSDL ngăn chặn tuyệt đối các nguy cơ tấn công chiếm quyền hoặc đọc trộm dữ liệu chéo giữa các dự án.")
    
    add_custom_heading(doc, "3.2.2 Hạn chế", level=3)
    add_bullet_paragraph(doc, "Hệ thống chưa tích hợp tính năng gửi thông báo tức thời (Real-time notifications) qua email hoặc đẩy trực tiếp trên giao diện khi có thay đổi trạng thái Story.")
    add_bullet_paragraph(doc, "Biểu đồ Burndown Chart hiện tại chưa hỗ trợ trừ đi các ngày nghỉ lễ, ngày cuối tuần trong chu kỳ Sprint, dẫn tới đường lý thuyết có thể bị chênh lệch nhẹ so với thực tế.")
    
    doc.add_page_break()
    
    # ------------------------------------------------------------------------------
    # KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN
    # ------------------------------------------------------------------------------
    p_conc_head = doc.add_paragraph()
    p_conc_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_conc_head.paragraph_format.space_before = Pt(12)
    p_conc_head.paragraph_format.space_after = Pt(18)
    run_conc_head = p_conc_head.add_run("KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN")
    set_run_font(run_conc_head, 'Times New Roman', 14, bold=True)
    
    add_custom_heading(doc, "Kết luận chung", level=2)
    add_body_paragraph(doc, "Đồ án cơ sở \"Xây dựng hệ thống quản lý dự án Scrum trực quan 'Kollab'\" đã hoàn thành đầy đủ các nội dung nghiên cứu lý thuyết và cài đặt thực tiễn theo đúng tiến độ đề ra. Hệ thống giải quyết tốt bài toán quản lý quy trình Agile/Scrum cho các đội ngũ phát triển phần mềm vừa và nhỏ thông qua các giao diện trực quan như Backlog, Kanban Board và các biểu đồ đo lường tiến độ tự động.")
    add_body_paragraph(doc, "Quá trình thực hiện đồ án đã giúp em củng cố sâu sắc kiến thức về lập trình web hiện đại (ReactJS, TypeScript, TailwindCSS), nắm vững mô hình Backend-as-a-Service với Supabase và học hỏi cách thiết lập cơ sở dữ liệu quan hệ PostgreSQL chuẩn hóa cũng như tư duy bảo mật dữ liệu nâng cao sử dụng Row Level Security.")
    
    add_custom_heading(doc, "Hướng phát triển tiếp theo", level=2)
    add_body_paragraph(doc, "Để đưa hệ thống Kollab trở thành một giải pháp thương mại hoàn chỉnh, các hướng nghiên cứu tiếp theo sẽ tập trung vào các điểm chính sau:")
    add_bullet_paragraph(doc, "Xây dựng hệ thống Real-time Notification tích hợp với Slack, Discord và Email để cập nhật tức thời các hành động của thành viên.")
    add_bullet_paragraph(doc, "Ứng dụng Trí tuệ nhân tạo (AI Assistant) để tự động hóa việc phân tích tốc độ hoàn thành Sprint (Velocity) và đưa ra các khuyến nghị tối ưu hóa phân bổ Story Points cho Scrum Master ở các Sprint kế tiếp.")
    add_bullet_paragraph(doc, "Tích hợp mô-đun quản lý tài liệu dự án trực tiếp (Wiki Project) để lưu trữ tài liệu đặc tả nghiệp vụ ngay bên trong ứng dụng Kollab.")
    
    doc.add_page_break()
    
    # ------------------------------------------------------------------------------
    # PHỤ LỤC (Đánh số trang riêng)
    # ------------------------------------------------------------------------------
    # Note: Appendix should start a new section with restarted page numbering,
    # but the template guidelines say "Phụ lục (Đánh số trang riêng)".
    # Let's create a new section for Appendix.
    sec_app = doc.add_section(start_type=WD_SECTION_START.NEW_PAGE)
    setup_section_margins(sec_app)
    sec_app.footer.is_linked_to_previous = False
    set_section_page_num_format(sec_app, "decimal", start=1)
    
    f_para_app = sec_app.footer.paragraphs[0]
    f_para_app.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_run_app = f_para_app.add_run()
    add_page_number_field(f_run_app)
    set_run_font(f_run_app, 'Times New Roman', 11)
    
    p_app_head = doc.add_paragraph()
    p_app_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_app_head.paragraph_format.space_before = Pt(12)
    p_app_head.paragraph_format.space_after = Pt(18)
    run_app_head = p_app_head.add_run("PHỤ LỤC")
    set_run_font(run_app_head, 'Times New Roman', 14, bold=True)
    
    add_custom_heading(doc, "Hướng dẫn cài đặt và cấu hình dự án Kollab", level=2)
    
    add_body_paragraph(doc, "Để khởi chạy hệ thống Kollab trên môi trường máy tính cá nhân (Local Development), vui lòng thực hiện tuần tự theo các bước hướng dẫn dưới đây:")
    
    add_body_paragraph(doc, "Bước 1: Chuẩn bị môi trường cài đặt", bold=True)
    add_bullet_paragraph(doc, "Đảm bảo máy tính đã cài đặt môi trường NodeJS phiên bản 18.0 trở lên (khuyên dùng bản LTS mới nhất).")
    add_bullet_paragraph(doc, "Đảm bảo đã cài đặt trình quản lý gói npm (thường đi kèm khi cài NodeJS) hoặc yarn.")
    
    add_body_paragraph(doc, "Bước 2: Clone mã nguồn và cài đặt dependencies", bold=True)
    p_code1 = doc.add_paragraph()
    p_code1.paragraph_format.left_indent = Cm(1.0)
    p_code1.paragraph_format.space_after = Pt(4)
    run_code1 = p_code1.add_run(
        "git clone https://github.com/user/kollab.git\n"
        "cd kollab\n"
        "npm install"
    )
    set_run_font(run_code1, 'Consolas', 10)
    
    add_body_paragraph(doc, "Bước 3: Cấu hình biến môi trường (Environment Variables)", bold=True)
    add_body_paragraph(doc, "Sao chép tệp biến môi trường mẫu và điền các thông tin kết nối tới dự án Supabase cá nhân của bạn:")
    p_code2 = doc.add_paragraph()
    p_code2.paragraph_format.left_indent = Cm(1.0)
    p_code2.paragraph_format.space_after = Pt(4)
    run_code2 = p_code2.add_run(
        "cp .env.example .env\n"
        "# Mở tệp .env và cập nhật thông số:\n"
        "VITE_SUPABASE_URL=https://your-project-id.supabase.co\n"
        "VITE_SUPABASE_ANON_KEY=your-anonymous-key-here"
    )
    set_run_font(run_code2, 'Consolas', 10)
    
    add_body_paragraph(doc, "Bước 4: Cấu hình cơ sở dữ liệu Supabase", bold=True)
    add_bullet_paragraph(doc, "Truy cập vào trang quản trị Supabase Dashboard, chọn dự án của bạn.")
    add_bullet_paragraph(doc, "Mở trình soạn thảo SQL (SQL Editor) và tạo mới một câu truy vấn.")
    add_bullet_paragraph(doc, "Sao chép toàn bộ mã lệnh SQL từ tệp supabase_schema.sql trong thư mục gốc dự án và dán vào SQL Editor của Supabase, sau đó nhấn Run (Chạy) để khởi tạo các bảng, triggers, indexes và các chính sách bảo mật RLS.")
    
    add_body_paragraph(doc, "Bước 5: Khởi chạy dự án ở môi trường phát triển", bold=True)
    p_code3 = doc.add_paragraph()
    p_code3.paragraph_format.left_indent = Cm(1.0)
    p_code3.paragraph_format.space_after = Pt(4)
    run_code3 = p_code3.add_run(
        "npm run dev"
    )
    set_run_font(run_code3, 'Consolas', 10)
    add_body_paragraph(doc, "Sau khi chạy lệnh trên, truy cập trình duyệt web theo đường dẫn http://localhost:5173 để bắt đầu trải nghiệm ứng dụng Kollab.")
    
    doc.add_page_break()
    
    # ------------------------------------------------------------------------------
    # DANH MỤC TÀI LIỆU THAM KHẢO
    # ------------------------------------------------------------------------------
    # Note: References are at the end, after appendix.
    # It says "đưa vào sau phụ lục, xếp theo thứ tự abc"
    p_ref_head = doc.add_paragraph()
    p_ref_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ref_head.paragraph_format.space_before = Pt(12)
    p_ref_head.paragraph_format.space_after = Pt(18)
    run_ref_head = p_ref_head.add_run("DANH MỤC TÀI LIỆU THAM KHẢO")
    set_run_font(run_ref_head, 'Times New Roman', 14, bold=True)
    
    ref_list = [
        "1. Ken Schwaber, Jeff Sutherland (2020), The Scrum Guide - The Definitive Guide to Scrum: The Rules of the Game, Scrum.Org (Link truy cập: https://scrumguides.org, truy cập ngày 10/05/2026).",
        "2. Nguyễn Văn An (2023), Lập trình ứng dụng Web hiện đại với ReactJS và TypeScript, Nhà xuất bản Thông tin và Truyền thông, Hà Nội.",
        "3. Nguyễn Văn Bằng (2024), Giáo trình Phân tích và Thiết kế hệ thống thông tin hướng đối tượng với UML, NXB Đại học Quốc gia TP. Hồ Chí Minh.",
        "4. Supabase Documentation (2026), Row Level Security in PostgreSQL, Supabase Inc. (Link truy cập: https://supabase.com/docs/guides/database/postgres/row-level-security, truy cập ngày 15/05/2026).",
        "5. Tailwind Labs (2025), TailwindCSS Utility-First Framework Documentation (Link truy cập: https://tailwindcss.com/docs, truy cập ngày 18/05/2026).",
        "6. Zustand Team (2025), Zustand State Management Guide for React (Link truy cập: https://zustand-demo.pmnd.rs, truy cập ngày 20/05/2026)."
    ]
    
    for r in ref_list:
        p_ref = doc.add_paragraph()
        p_ref.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_ref.paragraph_format.space_before = Pt(2)
        p_ref.paragraph_format.space_after = Pt(6)
        p_ref.paragraph_format.line_spacing = 1.5
        run_r = p_ref.add_run(r)
        set_run_font(run_r, 'Times New Roman', 13)
        
    # Save the document
    output_filename = "Bao_Cao_Do_An_Co_So_Kollab.docx"
    doc.save(output_filename)
    print(f"Document successfully created and saved as {output_filename}")

if __name__ == "__main__":
    create_report()
